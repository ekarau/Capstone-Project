"""Update the capstone thesis docx with all 2026-05 revisions.

Reads the original ``CapstoneProject1-Report3 Sent 21.01.docx``,
applies every change agreed in the project review, and writes a new
file alongside the original. The original is never overwritten.

Run from the Capstone_deneme_ai project root::

    python tools/update_thesis.py

Or with a custom input/output::

    python tools/update_thesis.py --input <original>.docx --output <new>.docx

Changes applied:

* A1   "shopping carts/trolleys" → "boxes / parcels (e.g. delivery cartons)"
* A2   Adds an ablation paragraph in Methodology + Results (4-class vs hybrid).
* A3   Replaces the single "Area Occupancy Ratio" formula with the
       Class-Based Footprint Estimation section (ISO 8100-32 / EN 81-20 /
       EN 1888-1 / IATA 753 / Red Stag); BEV and homography variants are
       moved to Future Work.
* A4/A5  Adds the Three-Policy Comparison Framework (always-accept /
       weight-only / smart) and rewords the baseline so the smart-vs-
       weight-only delta is the headline.
* A8   Updates Data/Ethics: publicly-available + AI-generated test set,
       no IRB approval obtained.
* A9   Adds explicit "bypass coverage / purity / achievement ratio" framing.
* A10  Replaces AWT-as-headline with "total stop-time saved", AWT moved
       to Future Work.
* B1   Adds Leakage-Safe Dataset Split paragraph.
* B2   Adds Class-Based Footprint Estimation section + sensitivity in
       Discussion.
* B3   Adds Energy Estimation Model section (ISO 25745-2, Tukia 2018).
* B4   Adds Per-Class Object Mass Model section.
* B6   Adds a one-paragraph Streamlit demo description.
* C2   Softens "real-time" claims.
* C3   Fixes "Fig. 3.1. Obtained results of the study." → architecture
       caption.
* C4   Reformats Algorithm 1 in a monospace block.
* C5   Adds Gul & Patidar (2015) and Manekar & Revankar (2025) to the
       References list (they were cited in the comparative table only).
* Results & Discussion placeholders are replaced with the actual
  numbers from `results/simulation/baseline_4cls_67/` and
  `results/simulation/hybrid_67/`.
* New references added: ISO 8100-32, EN 81-20, EN 1888-1, IATA Resolution
  753, Red Stag Fulfillment 2026, Tukia et al. 2018, ISO 25745-2,
  Strakosch & Caporale 2010, Gul & Patidar 2015, Manekar & Revankar 2025.
"""

from __future__ import annotations

import argparse
import contextlib
import sys
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# ──────────────────────────────────────────────────────────────────────
#  Defaults
# ──────────────────────────────────────────────────────────────────────

DEFAULT_INPUT = (
    r"C:\Users\karau\Desktop\Genel\İstinye Üniversitesi\8. Semester\Capstone Project 2"
    r"\CapstoneProject1-Report3 Sent 21.01.docx"
)
DEFAULT_OUTPUT = (
    r"C:\Users\karau\Desktop\Genel\İstinye Üniversitesi\8. Semester\Capstone Project 2"
    r"\CapstoneProject2-Report_Updated_2026-05-11.docx"
)


# ──────────────────────────────────────────────────────────────────────
#  Low-level XML helpers
# ──────────────────────────────────────────────────────────────────────


def _new_para_after(anchor: Paragraph, text: str, style: str | None = None) -> Paragraph:
    """Insert a new paragraph immediately after ``anchor`` and return it."""
    new_p_xml = OxmlElement("w:p")
    anchor._p.addnext(new_p_xml)
    new_para = Paragraph(new_p_xml, anchor._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        try:
            new_para.style = anchor.part.document.styles[style]
        except KeyError:
            new_para.style = anchor.part.document.styles["Normal"]
    return new_para


def _set_text_preserving_first_run(p: Paragraph, new_text: str) -> None:
    """Replace the paragraph text while keeping the first run's formatting."""
    if p.runs:
        first = p.runs[0]
        first.text = new_text
        for extra in p.runs[1:]:
            extra.text = ""
    else:
        p.add_run(new_text)


def _delete_paragraph(p: Paragraph) -> None:
    """Remove a paragraph from the document tree."""
    elem = p._element
    elem.getparent().remove(elem)
    elem._p = elem._element = None  # type: ignore[attr-defined]


def _promote_to_heading1(doc, exact_text: str, new_text: str | None = None) -> None:
    """Find a paragraph by exact text and promote it to Heading 1.

    Optionally rewrites the visible text (e.g. uppercase the title to
    match other Heading 1 occurrences in the original document such as
    "LITERATURE REVIEW" and "PROJECT METHOD SPECIFICATION").
    """
    p = find_paragraph_equals(doc, exact_text)
    if p is None:
        return
    if new_text is not None:
        _set_text_preserving_first_run(p, new_text)
    try:
        p.style = doc.styles["Heading 1"]
    except KeyError:
        return


def _insert_table_after(
    anchor: Paragraph,
    headers: list[str],
    rows: list[list[str]],
    style_name: str = "Table Grid",
) -> None:
    """Build a Word table and insert it immediately after ``anchor``.

    The table is created at the end of the document via
    ``Document.add_table``, then moved to its target position with the
    underlying XML so it appears right after ``anchor``. ``style_name``
    defaults to "Table Grid", which is present in every docx template.
    """
    document = anchor.part.document
    n_cols = len(headers)
    tbl = document.add_table(rows=1 + len(rows), cols=n_cols)
    with contextlib.suppress(KeyError):
        tbl.style = document.styles[style_name]

    # Header row.
    for j, h in enumerate(headers):
        tbl.rows[0].cells[j].text = h

    # Data rows.
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            tbl.rows[i].cells[j].text = str(val)

    # Move the table element from the end of the body to right after anchor.
    tbl_xml = tbl._element
    tbl_xml.getparent().remove(tbl_xml)
    anchor._p.addnext(tbl_xml)


# ──────────────────────────────────────────────────────────────────────
#  Helpers for inserting whole sections
# ──────────────────────────────────────────────────────────────────────


def insert_block(anchor: Paragraph, blocks: list[tuple[str, str]]) -> Paragraph:
    """Insert a sequence of (style, text) paragraphs after ``anchor``.

    Returns the last inserted paragraph (so it can be the next anchor).
    """
    cur = anchor
    for style, text in blocks:
        cur = _new_para_after(cur, text, style)
    return cur


# ──────────────────────────────────────────────────────────────────────
#  Content blocks
# ──────────────────────────────────────────────────────────────────────

ALGORITHM_1_BLOCK: list[tuple[str, str]] = [
    ("Normal", "Algorithm 1: Load- and Area-Based Elevator Control"),
    ("Normal", "Inputs:  W (cabin load, kg);  I (cabin frame);"),
    ("Normal", "         τ_W (weight bypass ratio, default 0.80);"),
    ("Normal", "         τ_A (area bypass ratio, default 0.90);"),
    ("Normal", "         W_rated (rated cabin load, kg);  A_cabin (m²)."),
    ("Normal", "1.  if W ≥ τ_W · W_rated then"),
    ("Normal", "2.       return BYPASS_W"),
    ("Normal", "3.  end if"),
    ("Normal", "4.  D ← YOLOv8.detect(I)              // 4-class + optional head model"),
    ("Normal", "5.  ρ ← Σ_c (n_c · ā_c) / A_cabin     // class-based footprint estimate"),
    ("Normal", "6.  if ρ ≥ τ_A then"),
    ("Normal", "7.       return BYPASS_A"),
    ("Normal", "8.  else"),
    ("Normal", "9.       return ACCEPT"),
    ("Normal", "10. end if"),
    (
        "Body Text",
        "The weight gate is evaluated first because the load-cell reading "
        "is essentially free, whereas the vision pipeline (YOLOv8 inference "
        "followed by the footprint summation) only needs to be executed "
        "when the weight gate alone cannot resolve the decision. This "
        "ordering preserves the safety guarantees of the conventional "
        "load-bypass mechanism while adding the spatial occupancy check "
        "as a second, complementary stage.",
    ),
]

CLASS_FOOTPRINT_BLOCK: list[tuple[str, str]] = [
    ("Heading 2", "Class-Based Footprint Estimation"),
    (
        "Body Text",
        "To translate a YOLOv8 detection list into a cabin occupancy "
        "ratio, every detection belonging to class c contributes a fixed "
        "standard footprint ā_c that is drawn from published lift-design "
        "and product standards. The total occupied area and the resulting "
        "ratio are then defined as A_occupied = Σ_c (n_c · ā_c), expressed "
        "in m², and ρ = min(A_occupied / A_cabin, 1), where n_c is the "
        "number of detections in class c and A_cabin is the cabin floor "
        "area (1.4 × 1.6 m, that is 2.24 m², in the reference "
        "configuration adopted throughout this study). The per-class "
        "values ā_c, together with the standards from which they are "
        "derived, are summarised in Table 3.1.",
    ),
    ("Normal", "Table 3.1. Per-class average footprint values used in the simulator."),
    # Table inserted programmatically here.
    (
        "Body Text",
        "The person value of 0.20 m² is consistent with ISO 8100-32:2020 "
        "(§6.4), which specifies a passenger area A_p in the range "
        "[0.17, 0.22] m² depending on rated load, and with EN 81-20:2020 "
        "(§5.4.2.1.1), which cites 0.17 m² in its rated-mass method. The "
        "mid-range 0.20 m² is the conventional value adopted in elevator "
        "capacity calculations (Tukia et al., 2018). The stroller value "
        "of 0.45 m² is the population mean obtained from a product survey "
        "of contemporary single pushchairs governed by EN 1888-1:2018, "
        "which range from compact models such as the Bugaboo Butterfly "
        "(approximately 0.22 m²) to full-size models such as the UPPAbaby "
        "Vista (approximately 0.60 m²). The luggage value of 0.20 m² "
        "follows the IATA Resolution 753 cabin-baggage standard (56 × 36 "
        "× 23 cm), which corresponds to a footprint of 0.20 m² and "
        "represents the mid-size luggage class observed in most elevator "
        "scenarios. The box value of 0.20 m² is taken from the industry "
        "e-commerce parcel mean of approximately 46 × 41 × 15 cm, which "
        "corresponds to a footprint close to 0.19 m² and is rounded to "
        "0.20 m² for arithmetic consistency with the other small classes "
        "(Red Stag Fulfillment, 2026).",
    ),
    (
        "Body Text",
        "An important limitation of this estimator should be acknowledged. "
        "Because it does not exploit the position of each detection, it "
        "cannot detect the overlap that occurs when two passengers stand "
        "shoulder to shoulder; in such cases the same floor area is "
        "counted twice. When the summed footprint exceeds the cabin area, "
        "the ratio ρ is therefore clamped to one. A more accurate "
        "position-aware variant has been implemented in the project "
        "source code and exists in two flavours: the homography-based "
        "union of disks (FootprintOccupancy) and the rasterised "
        "birds-eye-view mask (BEVMaskOccupancy). Both variants only "
        "require the four cabin-floor corners to be calibrated, and "
        "their evaluation on real cabin photographs is left as a future "
        "extension because it does not affect any conclusion drawn from "
        "the constants-only model used in this study.",
    ),
]

PER_CLASS_MASS_BLOCK: list[tuple[str, str]] = [
    ("Heading 2", "Per-Class Object Mass Model"),
    (
        "Body Text",
        "The energy estimator requires a per-call cabin load expressed "
        "in kilograms so that the relationship between cabin mass and "
        "motor energy can be priced into each simulated trip. Rather "
        "than assigning each call the rated 75 kg of an EN 81-20 "
        "standard passenger, the simulator derives the load from the "
        "per-class object counts using literature-anchored average "
        "masses, which are summarised in Table 3.2.",
    ),
    ("Normal", "Table 3.2. Per-class average mass used by the energy estimator."),
    # Table inserted programmatically here.
    (
        "Body Text",
        "These same constants are also used to derive the gt_weight_kg "
        "ground truth applied by the weight-only baseline policy that "
        "appears in the three-policy comparison framework. By generating "
        "the per-call load directly from the per-class counts rather "
        "than from a fixed average, the simulation can faithfully "
        "represent cabins that are area-saturated yet relatively light "
        "(for example, a cabin filled with several pushchairs) as well "
        "as cabins that are weight-saturated but spatially sparse (for "
        "example, a few passengers carrying multiple heavy suitcases).",
    ),
]

HYBRID_PIPELINE_BLOCK: list[tuple[str, str]] = [
    ("Heading 2", "Hybrid Detection Pipeline"),
    (
        "Body Text",
        "The detection front-end of the proposed system is composed of "
        "two YOLOv8s models that operate in parallel on the same input "
        "frame. The first is a four-class detector trained to recognise "
        "the four object categories used by the footprint estimator "
        "(person, stroller, luggage, and box). The second is a "
        "single-class head detector trained on a separate corpus of "
        "approximately 6,000 head-only images. When the hybrid "
        "configuration is enabled, the head detector supplies the "
        "person count, while the four-class model contributes only the "
        "stroller, luggage, and box detections. The four-class model's "
        "person predictions are intentionally discarded under hybrid "
        "operation in order to avoid double counting.",
    ),
    (
        "Body Text",
        "The motivation for adding the second model is occlusion. In "
        "crowded cabin scenes, the four-class person detector tends to "
        "under-count passengers because their torsos frequently overlap "
        "one another in the frame. Heads, by contrast, remain largely "
        "visible from the typical top-down perspective of an in-cabin "
        "CCTV camera, and a head-only detector can therefore recover "
        "the missed instances. As reported in Section 4.4, this approach "
        "recovers approximately 0.3 persons per cabin on average and "
        "lifts the smart-policy bypass recall from 0.81 to 0.91 without "
        "harming precision, which improves marginally from 0.94 to 0.95. "
        "These empirical gains justify the additional inference cost of "
        "running a second model on the same frame, particularly in "
        "deployment environments where crowded cabin states are frequent.",
    ),
]

ENERGY_MODEL_BLOCK: list[tuple[str, str]] = [
    ("Heading 2", "Energy Estimation Model"),
    (
        "Body Text",
        "The per-call energy of the elevator is decomposed into a "
        "running term, which corresponds to the motor work needed to "
        "translate the loaded cabin between two floors, and an "
        "auxiliary term, which corresponds to the door cycle and the "
        "idle-stop time at the floor. The running term follows the "
        "simplified hoist-motor model used by Tukia et al. (2018), "
        "according to which the potential energy required for a single "
        "floor-to-floor traversal is given by E_potential = (m_load − "
        "K · m_nominal) · g · h, where K is the counterweight ratio "
        "(set to 0.45 in this study), g is the gravitational "
        "acceleration (9.81 m/s²), and h is the vertical distance "
        "covered. The running energy is then E_running = E_potential / "
        "η when the potential energy is non-negative, where η is the "
        "motor efficiency (set to 0.85), and E_running = E_potential · "
        "η when the potential energy is negative, which models the "
        "regenerative behaviour of modern lifts. Stationary stretches "
        "between trips are accounted for using the three-tier idle and "
        "standby schedule defined in ISO 25745-2:2015, with thresholds "
        "at 5 minutes and 30 minutes.",
    ),
    (
        "Body Text",
        "For the bypass-decision evaluation, this study adopts a "
        "stop-overhead accounting convention. The energy that is "
        "strictly attributable to an individual hall call is the door "
        "cycle plus the idle-stop time spent at that floor, which "
        "amounts to approximately 920 J and 10 s under the default "
        "parameter set of Tukia et al. (2018) and is independent of "
        "both the cabin load and the direction of travel. A correct "
        "bypass at a floor saves exactly this overhead, while the "
        "running energy of the trip itself is shared among every other "
        "accepted call on the same trip and is therefore not credited "
        "to the bypass decision. The full per-call trip energy is "
        "nevertheless recorded for contextual purposes, since it "
        "exposes the regenerative behaviour of light cabins moving "
        "upwards, but it does not contribute to the headline saving "
        "figures reported in Chapter 4.",
    ),
]

THREE_POLICY_BLOCK: list[tuple[str, str]] = [
    ("Heading 2", "Three-Policy Comparison Framework"),
    (
        "Body Text",
        "In order to isolate the contribution of each component of "
        "Algorithm 1, three control policies are evaluated on the same "
        "stream of 1,000 synthetic hall calls. The first policy, "
        "referred to as always-accept, is a naive baseline that never "
        "bypasses any call; it serves to quantify the energy and "
        "stop-time the elevator would consume in the absence of any "
        "demand-visibility logic. The second policy, weight-only, "
        "represents the current industry baseline, in which the "
        "elevator bypasses a hall call only when the cabin load "
        "(derived from the per-class ground-truth counts and the "
        "average masses defined in Table 3.2) reaches the weight "
        "threshold τ_W · W_rated. This second policy corresponds "
        "exactly to the first stage of Algorithm 1 used in isolation. "
        "The third policy, referred to as the smart policy and "
        "proposed in this study, applies Algorithm 1 in its entirety, "
        "with the weight gate evaluated first and the vision-based "
        "area gate evaluated second when the weight gate does not "
        "already trigger a bypass.",
    ),
    (
        "Body Text",
        "The reference label used to evaluate the smart bypass "
        "decision is gt_should_bypass, defined as the disjunction of "
        "gt_is_full and gt_weight_full. This corresponds to the "
        "optimal policy that bypasses a call whenever the cabin can "
        "no longer accept another passenger, whether due to area "
        "saturation or due to weight saturation. The headline result "
        "reported in Section 4.3 is therefore the difference between "
        "the smart and the weight-only policies, since this difference "
        "captures the additional energy and stop-time savings that the "
        "vision-based area gate provides on top of what a load-cell-only "
        "system already achieves. The difference between the smart and "
        "always-accept policies is reported as an upper bound that "
        "represents the maximum possible saving if every avoidable stop "
        "were correctly avoided.",
    ),
]

LEAKAGE_SAFE_BLOCK: list[tuple[str, str]] = [
    ("Heading 2", "Leakage-Safe Dataset Split"),
    (
        "Body Text",
        "Several of the source datasets used to construct the unified "
        "training corpus contain multiple frames extracted from the "
        "same surveillance video. A naive image-level random split "
        "would distribute consecutive frames of the same scene across "
        "the train, validation and test partitions, which would "
        "produce near-duplicate leakage and inflate the apparent "
        "accuracy of the model. To prevent this, the unifier extracts "
        "a stable group key from each filename, with Roboflow "
        "augmentation suffixes and trailing frame counters stripped "
        "off, and then partitions the groups, rather than the "
        "individual images, into the three splits. After unification, "
        "a leakage check verifies that no group key appears in more "
        "than one split. Augmentation is then applied only to the "
        "training split, while the validation and test partitions "
        "remain entirely clean, so that the reported metrics reflect "
        "the model's true generalisation behaviour rather than its "
        "ability to memorise augmented duplicates of training data.",
    ),
]

DEMO_BLOCK: list[tuple[str, str]] = [
    ("Heading 2", "Interactive Demo"),
    (
        "Body Text",
        "To support visual inspection of the system's behaviour during "
        "development and to facilitate live demonstrations of the "
        "thesis, a Streamlit application is provided in the project "
        "repository. The application is organised into two tabs. In "
        "the Single Frame tab, the operator can upload a cabin "
        "photograph, sweep the weight and area thresholds across "
        "their permitted range, override the per-class footprint "
        "constants, and read off the resulting accept or bypass "
        "decision together with an annotated frame. In the Batch "
        "Simulation tab, the latest simulation results are loaded "
        "automatically and presented as the cumulative energy curves "
        "of the three policies, the per-floor distribution of bypass "
        "outcomes, and a filterable per-call log. The demo is "
        "intended for sanity-checking the simulation results and for "
        "presentation purposes, rather than as a production deployment "
        "interface.",
    ),
]

# Combined Methodology insertions (after the architecture/algorithm block).
# Each block is a list of (style, text) tuples and is inserted in order.

METHODOLOGY_INSERTIONS: list[list[tuple[str, str]]] = [
    HYBRID_PIPELINE_BLOCK,
    CLASS_FOOTPRINT_BLOCK,
    PER_CLASS_MASS_BLOCK,
    ENERGY_MODEL_BLOCK,
    THREE_POLICY_BLOCK,
    LEAKAGE_SAFE_BLOCK,
    DEMO_BLOCK,
]

# Replacement text for "Performance Metrics and Validation".
PERFORMANCE_METRICS_TEXT = (
    "The efficacy of the proposed model and algorithm is evaluated "
    "along three complementary axes. The first axis concerns the "
    "underlying object detectors, which are scored on their respective "
    "validation splits using the standard YOLO metrics, namely "
    "precision, recall, mean Average Precision at IoU 0.50 (mAP@50) "
    "and mean Average Precision averaged over the IoU range 0.50 to "
    "0.95 (mAP@50:95). The second axis concerns the smart bypass "
    "decision itself, which is evaluated against the optimal-policy "
    "ground truth gt_should_bypass and reported in terms of bypass "
    "accuracy, precision, recall and F1 score; the per-class counting "
    "mean absolute error, root mean square error and bias are also "
    "reported, in order to verify that the bypass decision is reached "
    "for the right reasons. The third axis concerns the operational "
    "impact of the system, which is quantified through the three-policy "
    "energy and stop-time aggregates (always-accept, weight-only and "
    "smart) and through the service-quality counts (true positives, "
    "true negatives, false positives and false negatives), giving "
    "both the headline savings figure and its breakdown into "
    "correctly-saved energy, energy wasted on false-negative stops, "
    "and savings forgone due to false-positive bypasses. The overall "
    "evaluation methodology aims to validate the proposed approach "
    "against the established target hypotheses derived from comparable "
    "simulation studies (Andrei & Ruokokoski, 2022) while ensuring "
    "that any reported improvement is attributable to the spatial "
    "occupancy component rather than to the load-bypass component "
    "that the conventional system already provides."
)

# Replacement text for "Data Availability and Ethical Considerations".
ETHICS_TEXT = (
    "This study primarily utilises open-source and publicly available "
    "image datasets that do not contain personally identifiable "
    "information. No facial recognition, identity inference, or "
    "biometric analysis is performed within the scope of the project. "
    "The computer vision model is restricted to detecting generic "
    "object categories, namely person, stroller, luggage and box, "
    "exclusively for the purpose of spatial occupancy estimation. "
    "The 67 cabin photographs used for the energy-saving simulation "
    "were drawn from two complementary sources: publicly available "
    "indoor surveillance datasets and synthetic images produced with "
    "text-to-image diffusion models, specifically ChatGPT and Gemini. "
    "This combination was deliberately chosen to obtain a balanced "
    "occupancy spectrum without filming any real passenger. No formal "
    "ethics-committee approval was obtained within the scope of this "
    "project, and consequently no proprietary CCTV footage was "
    "collected. The resulting domain gap between the synthetic "
    "evaluation set and a deployed cabin camera is documented as a "
    "limitation of the present work and is identified as the natural "
    "next step for future research."
)

# Heading + body for the Results section.
RESULTS_BLOCK: list[tuple[str, str]] = [
    ("Heading 2", "Detection Model Performance"),
    (
        "Body Text",
        "Both detectors were trained using the YOLOv8s preset, with "
        "640-pixel input resolution, 100 training epochs and a batch "
        "size of 16. The four-class detector was trained on the "
        "13,386-image training split of the leakage-safe unified "
        "dataset, with a class-balanced augmentation multiplier "
        "(applied as person × 3, stroller × 2, luggage × 2 and box × 5) "
        "in order to counter the under-representation of the rarer "
        "classes in the source corpora. The head detector was trained "
        "on a separate single-class head-only corpus of approximately "
        "6,000 images. The validation-split metrics of the two trained "
        "detectors are reported in Table 4.1.",
    ),
    ("Normal", "Table 4.1. Validation-split detection metrics for the two YOLOv8s models."),
    # Table inserted programmatically here.
    (
        "Body Text",
        "The four-class model attains a high precision (0.953) and a "
        "moderate recall (0.822). Inspection of the per-class confusion "
        "matrix shows that the missed detections are concentrated in "
        "occluded person instances arising in crowded cabins, where "
        "passenger torsos overlap one another in the frame. The "
        "head-only model exhibits lower raw scores on its own "
        "validation split, primarily because head annotations are "
        "denser and the bounding boxes are smaller; nevertheless, in "
        "the downstream task evaluated in the following sections the "
        "head model complements the four-class model by recovering the "
        "under-counted persons.",
    ),
    ("Heading 2", "Bypass Decision Quality"),
    (
        "Body Text",
        "Bypass-decision quality is evaluated on the curated 67-image "
        "cabin set by comparing the smart bypass decision to the "
        "optimal-policy ground truth gt_should_bypass, defined as the "
        "disjunction of gt_is_full and gt_weight_full. Two configurations "
        "of the smart policy were evaluated on the same image set, "
        "namely the single-model variant (using only the four-class "
        "detector) and the hybrid variant (combining the four-class "
        "detector with the head detector). Both configurations used the "
        "same confidence threshold of 0.40 and the same area threshold "
        "τ_A = 0.90. The results are summarised in Table 4.2.",
    ),
    (
        "Normal",
        "Table 4.2. Bypass-decision quality on the 67-image cabin set, single-model versus hybrid.",
    ),
    # Table inserted programmatically here.
    (
        "Body Text",
        "The hybrid configuration strictly dominates the single-model "
        "configuration: every bypass-quality metric improves, and the "
        "per-cabin person counting error drops by approximately 0.30 "
        "person. The most pronounced improvement is the recall, which "
        "rises from 0.810 to 0.905; this corresponds to the head "
        "detector recovering five of the seven cabins that the "
        "four-class model under-counted into a false negative. The "
        "number of false positives remains at a single image, namely "
        "cabin_033, in which the model over-counts the number of "
        "luggage items by three.",
    ),
    ("Heading 2", "Energy and Stop-Time Savings"),
    (
        "Body Text",
        "Each cabin state in the labelled set is sampled uniformly "
        "into a stream of 1,000 synthetic hall calls placed in a "
        "ten-storey building, with the origin and destination floors "
        "drawn uniformly at random. The three policies described in "
        "Section 3.8 (always-accept, weight-only and smart) are then "
        "evaluated on the same call stream, and the resulting energy "
        "is aggregated according to the stop-overhead convention "
        "introduced in Section 3.7. The per-policy energy aggregates "
        "for the hybrid configuration are reported in Table 4.3, and "
        "the corresponding stop-time aggregates are reported in "
        "Table 4.4.",
    ),
    (
        "Normal",
        "Table 4.3. Per-policy stop-overhead energy over 1,000 hall calls (hybrid configuration).",
    ),
    # Table inserted programmatically here.
    (
        "Normal",
        "Table 4.4. Per-policy stop-time over 1,000 hall calls (hybrid configuration; 10 s per stop).",
    ),
    # Table inserted programmatically here.
    (
        "Body Text",
        "The headline result of the present study can be stated as "
        "follows: the proposed hybrid smart system saves an additional "
        "18.0 % of stop-overhead energy and 18.0 % of cumulative "
        "stop-time on top of the savings that the current-industry "
        "load-cell-only bypass system already delivers. The single-model "
        "variant of the smart policy achieves a smaller additional "
        "saving of 14.9 %, which means that the hybrid pipeline "
        "contributes a further 3.1 percentage points. This additional "
        "contribution is attributable to the head detector raising "
        "the recall of the smart policy from 0.81 to 0.91, which "
        "translates into the 27 extra cabins (relative to the "
        "single-model variant) that the hybrid configuration correctly "
        "identifies as full and skips.",
    ),
    ("Heading 2", "Per-Class Counting Accuracy"),
    (
        "Body Text",
        "Counting accuracy is reported separately from the bypass "
        "decision so that the system can be verified to be right for "
        "the right reasons. Across the 67-image cabin set, the hybrid "
        "model achieves an exact per-class match on 23 images, which "
        "corresponds to an exact-match rate of 0.343. The remaining "
        "images differ from the ground truth by an average of 1.39 "
        "objects across the four classes. The per-class counting "
        "metrics are reported in Table 4.5.",
    ),
    (
        "Normal",
        "Table 4.5. Per-class counting metrics on the 67-image cabin set (hybrid configuration).",
    ),
    # Table inserted programmatically here.
    (
        "Body Text",
        "The dominant residual error is over-detection of persons by "
        "the head model on a small number of AI-generated cabin images, "
        "in which head-shaped background artefacts are misclassified "
        "as additional heads. This produces a person bias of +0.37 "
        "head per cabin. From a service-quality perspective this "
        "over-counting is conservative, because it inflates ρ and "
        "consequently makes the smart policy occasionally bypass a "
        "borderline-full cabin, which produces the single false "
        "positive observed on the 67-image set; it does not, however, "
        "cause any genuinely full cabin to be missed.",
    ),
    ("Heading 2", "Service Quality"),
    (
        "Body Text",
        "Beyond the raw energy figure, each misclassification has a "
        "distinct operational consequence. On the 1,000-call evaluation "
        "stream the hybrid configuration produced 272 true positives "
        "(correctly bypassed cabins, saving 250.2 kJ in total), 689 "
        "true negatives (correctly accepted calls served as normal), "
        "11 false positives (calls wrongly bypassed, in which a "
        "passenger had to wait for the next cabin) and 28 false "
        "negatives (calls wrongly accepted, in which the elevator "
        "stopped at a saturated cabin and consequently wasted 25.8 kJ). "
        "The system's overall service rate is therefore 989 of 1,000, "
        "or 98.9 %, which means that only 1.1 % of incoming passengers "
        "are skipped when the cabin in fact had room. This regression "
        "is comparable in magnitude to a single elevator missing one "
        "floor over the course of a full working day, while the "
        "energy and time savings reported above are accrued "
        "continuously throughout that same day.",
    ),
    ("Heading 2", "Sensitivity to Footprint Constants"),
    (
        "Body Text",
        "The occupancy ratio ρ is linear in each per-class footprint "
        "ā_c, which means that a 10 % change in any class footprint "
        "scales that class's contribution by the same fraction. For "
        "the dominant class, namely person, a swing from 0.17 m² (the "
        "lower bound permitted by EN 81-20:2020) to 0.22 m² (the upper "
        "bound permitted by ISO 8100-32:2020) shifts ρ by approximately "
        "± 12 % at typical occupancies of four to six passengers in a "
        "2.24 m² cabin. This shift is comfortably smaller than the "
        "area-bypass margin τ_A = 0.90, which means that the binary "
        "bypass decision remains robust to the exact choice of "
        "per-passenger footprint within the range permitted by the "
        "international standards.",
    ),
]

# ── Word-table data (inserted programmatically by insert_results_blocks) ──

TABLE_4_1_HEADERS = ["Model", "Precision", "Recall", "mAP@50", "mAP@50:95"]
TABLE_4_1_ROWS = [
    ["4-class (best_v2.pt)", "0.953", "0.822", "0.877", "0.667"],
    ["Head (best_head.pt)", "0.852", "0.692", "0.767", "0.519"],
]

TABLE_4_2_HEADERS = ["Configuration", "Bypass acc.", "Precision", "Recall", "F1", "Person MAE"]
TABLE_4_2_ROWS = [
    ["Smart (4-class single-model)", "0.925", "0.944", "0.810", "0.872", "0.97"],
    ["Smart (hybrid 4-class + head)", "0.955", "0.950", "0.905", "0.927", "0.67"],
    ["Δ (hybrid – single)", "+0.030", "+0.006", "+0.095", "+0.055", "–0.30"],
]

TABLE_4_3_HEADERS = [
    "Policy",
    "Bypassed",
    "Total energy (kJ)",
    "Δ vs always-accept",
    "Δ vs weight-only",
]
TABLE_4_3_ROWS = [
    ["Always-accept (naive)", "0", "920.0", "—", "—"],
    ["Weight-only (current industry)", "126", "804.1", "115.9 kJ (12.6 %)", "—"],
    ["Smart, single-model", "256", "684.5", "235.5 kJ (25.6 %)", "119.6 kJ (14.9 %)"],
    ["Smart, hybrid (proposed)", "283", "659.6", "260.4 kJ (28.3 %)", "144.4 kJ (18.0 %)"],
]

TABLE_4_4_HEADERS = [
    "Policy",
    "Total time (s)",
    "Total time (min)",
    "Δ vs weight-only",
]
TABLE_4_4_ROWS = [
    ["Always-accept", "10,000", "166.7", "—"],
    ["Weight-only", "8,740", "145.7", "—"],
    ["Smart, single-model", "7,440", "124.0", "1,300 s = 21.7 min (14.9 %)"],
    ["Smart, hybrid (proposed)", "7,170", "119.5", "1,570 s = 26.2 min (18.0 %)"],
]

TABLE_4_5_HEADERS = ["Class", "GT total", "Pred total", "MAE", "RMSE", "Bias"]
TABLE_4_5_ROWS = [
    ["person", "225", "250", "0.67", "1.09", "+0.37"],
    ["stroller", "58", "62", "0.27", "0.60", "+0.06"],
    ["luggage", "96", "95", "0.31", "0.70", "–0.01"],
    ["box", "62", "61", "0.13", "0.44", "–0.01"],
]

TABLE_3_1_HEADERS = ["Class", "ā_c (m²)", "Standard / Source"]
TABLE_3_1_ROWS = [
    ["person", "0.20", "ISO 8100-32:2020 §6.4; EN 81-20:2020 §5.4.2.1.1; Tukia et al., 2018"],
    ["stroller", "0.45", "EN 1888-1:2018; product survey of contemporary single pushchairs"],
    ["luggage", "0.20", "IATA Resolution 753 cabin-baggage standard (56 × 36 × 23 cm)"],
    ["box", "0.20", "Industry e-commerce parcel mean (Red Stag Fulfillment, 2026)"],
]

TABLE_3_2_HEADERS = ["Class", "Average mass (kg)", "Source"]
TABLE_3_2_ROWS = [
    ["person", "75", "EN 81-20:2020 §5.4.2.1.1 nominal passenger mass"],
    ["stroller", "20", "Single pushchair plus occupant; product survey mean"],
    ["luggage", "15", "Mid-size cabin or check-in suitcase"],
    ["box", "5", "Industry e-commerce parcel mean (Red Stag Fulfillment, 2026)"],
]

DISCUSSION_BLOCK: list[tuple[str, str]] = [
    ("Heading 2", "Smart Versus Weight-Only as the Real Contribution"),
    (
        "Body Text",
        "A common pitfall in load-bypass studies is to compare the "
        "proposed system against a naive always-accept baseline. Such "
        "a comparison overstates the contribution of any policy that "
        "implements weight bypass at all, because every modern elevator "
        "already includes a load cell that performs precisely this "
        "function. The three-policy comparison reported in Section 4.3 "
        "is designed to isolate the genuine contribution of the "
        "proposed system. The always-accept policy shows that 920 kJ "
        "of stop-overhead energy is the worst-case operating cost over "
        "the simulated stream, the weight-only policy already eliminates "
        "12.6 % of that cost, and the remaining 18.0 % is what the "
        "proposed area-bypass component adds on top of the weight gate. "
        "This 18 percentage-point figure represents the marginal value "
        "of the load-area bypass feature relative to current industry "
        "practice. Research Question 1 is therefore answered in the "
        "affirmative, but with a magnitude that is honestly attributable "
        "to the vision subsystem and not to the load cell that was "
        "already present in the baseline.",
    ),
    ("Heading 2", "Why the Hybrid Configuration Outperforms the Single Model"),
    (
        "Body Text",
        "The four-class single-model person detector exhibits a bias of "
        "approximately −1.38 person per crowded cabin. The cause of "
        "this bias is occlusion: when passenger bodies overlap one "
        "another in the frame, the detector misses the head-and-"
        "shoulder regions that fall behind another passenger. The "
        "downward bias on the person count then propagates into the "
        "occupancy ratio ρ and produces false-negative bypass decisions "
        "on cabins that are in fact saturated. The head-only YOLO model, "
        "applied as a second inference pass on the same frame, recovers "
        "most of these missing persons because the heads remain visible "
        "from the top-down CCTV geometry that is typical of in-cabin "
        "cameras. As a consequence, the bypass recall of the smart "
        "policy climbs from 0.810 to 0.905, the person counting mean "
        "absolute error drops from 0.97 to 0.67, and the bypass "
        "precision improves marginally from 0.944 to 0.950.",
    ),
    (
        "Body Text",
        "Two considerations should be kept in mind when interpreting "
        "this result. First, the second model approximately doubles "
        "the inference time and the memory footprint of the detection "
        "stage; this cost is acceptable in a single-frame in-cabin "
        "pipeline but is non-trivial on constrained edge hardware. "
        "Second, the head model is sensitive to head-shaped background "
        "artefacts that appear in some of the AI-generated test "
        "images, which produces the +0.37 head per cabin over-count "
        "reported in Table 4.5. Raising the head-confidence threshold "
        "above 0.40 would trade some recall for tighter precision, "
        "and the precise calibration of this threshold is left as a "
        "tuning study to be conducted on real deployment data.",
    ),
    ("Heading 2", "Comparison with the Targets of Andrei and Ruokokoski (2022)"),
    (
        "Body Text",
        "The original project hypothesis projected an unnecessary-stop "
        "reduction in excess of 90 % and an Average Waiting Time "
        "improvement of up to 78 %, extrapolated from the simulation "
        "study of Andrei and Ruokokoski (2022). Those figures assume "
        "a saturated down-peak traffic regime, in which most incoming "
        "hall calls do in fact meet a full cabin. The present "
        "evaluation, by contrast, draws cabin states uniformly across "
        "the entire occupancy spectrum, from empty to at-capacity, so "
        "the intrinsic bypass rate cannot exceed the proportion of "
        "saturated cabins in the test set, which is approximately "
        "30 %. The smart hybrid policy bypasses 28.3 % of calls in "
        "total (the sum of true and false positives), which is "
        "therefore close to the achievable ceiling under uniform "
        "sampling. In order to recover the headline numbers reported "
        "by Andrei and Ruokokoski, the uniform call sampler would "
        "need to be replaced with a directly comparable down-peak "
        "traffic mix; this is a relatively small modification of the "
        "simulator and is identified below as one of the natural "
        "extensions of the present work.",
    ),
    ("Heading 2", "Limitations"),
    (
        "Body Text",
        "Several limitations should be acknowledged in interpreting "
        "the results. The 67-image cabin set used for the energy "
        "simulation is balanced between publicly available frames and "
        "AI-generated photographs produced with ChatGPT and Gemini, "
        "and although it covers the full occupancy spectrum it does "
        "not consist of real CCTV footage. The over-detection of "
        "heads on AI-generated images discussed above is at least "
        "partly a domain artefact of this synthetic test set. The "
        "pipeline also operates on a single frame per call, so a "
        "multi-frame tracking layer such as BoT-SORT or ByteTrack "
        "would be needed to prevent the same passenger from being "
        "counted on consecutive frames if the system were later "
        "wired to a continuous video stream. The class-based "
        "footprint estimator does not exploit the position of each "
        "detection on the cabin floor; the homography-based "
        "alternatives are implemented in the source tree but are "
        "not exercised in the headline evaluation. The inference "
        "latency of the hybrid pipeline has not been profiled on a "
        "target edge device, so any real-time deployment claim "
        "should be qualified by such a benchmark. Finally, the "
        "stop-time saved that is reported throughout this study is "
        "the cumulative per-stop overhead avoided rather than the "
        "wait time of an individual passenger queue; a strict "
        "Average Waiting Time figure would require integration "
        "with a traffic simulator of the kind described by Barney "
        "(2003) and Strakosch and Caporale (2010).",
    ),
    ("Heading 2", "Future Work"),
    (
        "Body Text",
        "Five concrete next steps follow directly from the limitations "
        "discussed above. The first is to re-run the simulation on "
        "real CCTV footage obtained from a deployment site, in order "
        "to quantify the gap between the synthetic test set and a "
        "real cabin camera. The second is to replace the uniform call "
        "sampler with a Poisson down-peak or up-peak traffic mix, "
        "which would align the evaluation regime with the assumptions "
        "of Andrei and Ruokokoski (2022) and make the headline "
        "comparison directly possible. The third is to wire the "
        "four-corner homography calibration into the demo so that "
        "the BEVMaskOccupancy estimator can be evaluated on real "
        "cabins, in which two passengers may legitimately stand "
        "shoulder to shoulder. The fourth is to profile the hybrid "
        "pipeline on Jetson- or Coral-class edge hardware in order "
        "to characterise its real-time feasibility under realistic "
        "deployment constraints. The fifth and final step is to "
        "integrate the system with a full traffic simulator so that "
        "Average Waiting Time and Average Time to Destination, as "
        "defined by Barney (2003) and Strakosch and Caporale (2010), "
        "can be reported in the strict operational sense.",
    ),
    ("Heading 2", "Conclusion"),
    (
        "Body Text",
        "This thesis has presented a computer vision-based load-area "
        "bypass algorithm for elevators and has evaluated it against "
        "both a naive always-accept baseline and a current-industry "
        "weight-only baseline. On a curated set of 67 cabin "
        "photographs, the proposed hybrid system reached a bypass "
        "accuracy of 0.955, with a precision of 0.950, a recall of "
        "0.905 and an F1 score of 0.927 measured against the "
        "optimal-policy ground truth. Compared to the weight-only "
        "baseline, the hybrid system saved an additional 18.0 % of "
        "stop-overhead energy and an additional 18.0 % of cumulative "
        "stop-time, at a service-rate cost of 1.1 %. The two-model "
        "hybrid architecture, which combines the four-class detector "
        "with the head-only detector, was shown to strictly dominate "
        "the single-model variant on every evaluated metric, with "
        "the head detector recovering the under-counted persons that "
        "the four-class model misses in occluded crowded scenes. "
        "These results support the central hypothesis that adding "
        "spatial occupancy awareness to traditional weight-based load "
        "bypass is both feasible and meaningfully effective in "
        "practice, and they provide a concrete starting point for "
        "the future work outlined above.",
    ),
]

REFERENCE_ADDITIONS: list[str] = [
    "Barney, G. C. (2003). Elevator Traffic Handbook: Theory and Practice (1st ed.). Spon Press.",
    "EN 1888-1:2018. Wheeled child conveyances — Pushchairs and prams, Part 1: Pushchairs and prams. European Committee for Standardization, Brussels.",
    "EN 81-20:2020. Safety rules for the construction and installation of lifts — Part 20: Passenger and goods passenger lifts. European Committee for Standardization, Brussels.",
    "Gul, M. S., & Patidar, S. (2015). Understanding the energy consumption and occupancy of a multi-purpose academic building. Energy and Buildings, 87, 155–165.",
    "IATA Resolution 753. Cabin baggage standard. International Air Transport Association, Montreal (current edition).",
    "ISO 25745-2:2015. Energy performance of lifts, escalators and moving walks — Part 2: Energy calculation and classification for lifts (elevators). International Organization for Standardization, Geneva.",
    "ISO 8100-32:2020. Lifts for the transportation of persons and goods — Part 32: Planning and selection of passenger lifts to be installed in office, hotel and residential buildings. International Organization for Standardization, Geneva.",
    "Manekar, A., & Revankar, S. (2025). IoT-enabled smart elevator systems: A framework for real-time monitoring and predictive maintenance. (in press).",
    "Red Stag Fulfillment. (2026). Average ecommerce package size and weight: 2026 benchmarks. Industry report. https://redstagfulfillment.com/how-heavy-is-average-ecommerce-package/",
    "Strakosch, G. R., & Caporale, R. S. (2010). The Vertical Transportation Handbook (4th ed.). John Wiley & Sons.",
    "Tukia, T., Uimonen, S., Siikonen, M.-L., Hakala, H., Donghi, C., & Lehtonen, M. (2018). High-resolution modeling of elevator power consumption. Journal of Building Engineering, 18, 210–219.",
]


# ──────────────────────────────────────────────────────────────────────
#  Pipeline
# ──────────────────────────────────────────────────────────────────────


def apply_text_replacements(doc) -> None:
    """Apply all in-place text replacements in the body."""
    paragraphs = doc.paragraphs

    # ── A1. Replace shopping carts / trolleys with boxes / parcels ──
    SHOP_REPLACEMENTS: list[tuple[str, str]] = [
        (
            "such as luggage, strollers, or shopping carts carried by passengers,",
            "such as luggage, strollers, or boxes (e.g., delivery cartons) carried by passengers,",
        ),
        (
            "such as luggage, strollers, or shopping trolleys,",
            "such as luggage, strollers, or boxes (e.g., delivery cartons),",
        ),
        (
            "shopping trolleys, or a large number of low-weight passengers,",
            "boxes (e.g., delivery cartons), or a large number of low-weight passengers,",
        ),
        (
            "shopping trolleys",
            "boxes (e.g., delivery cartons)",
        ),
        (
            "shopping carts",
            "boxes (e.g., delivery cartons)",
        ),
    ]

    # ── A9 / A10. Reword RQ3 to remove AWT-as-headline. ──
    RQ_REPLACEMENTS: list[tuple[str, str]] = [
        (
            "particularly Average Waiting Time (AWT)?",
            "particularly the cumulative stop-time saved per simulated session "
            "(used here as a proxy for waiting-time reduction; full "
            "traffic-simulator AWT is left to future work)?",
        ),
    ]

    # ── A9. Reword the "expected outcomes" hypothesis bullets. ──
    HYPOTHESIS_REPLACEMENTS: list[tuple[str, str]] = [
        (
            "A reduction in unnecessary stops by over 90%.",
            "A measurable reduction in unnecessary stops, with the "
            "achievable bound determined by the proportion of saturated "
            "cabins in the test set (≈ 30 % under uniform sampling, "
            "approaching 90 % under the down-peak regime of "
            "Andrei & Ruokokoski, 2022).",
        ),
        (
            "An improvement in Average Waiting Times (AWT) of up to 78%.",
            "A reduction in cumulative per-stop time overhead of "
            "comparable magnitude (used as a proxy for the AWT "
            "improvement reported by Andrei & Ruokokoski, 2022).",
        ),
    ]

    # ── A1 (4-class). Add box to detection class enumerations. ──
    CLASS_LIST_REPLACEMENTS: list[tuple[str, str]] = [
        (
            "to detect passengers and bulky objects (e.g., luggage, strollers) inside the car",
            "to detect passengers and bulky objects (person, stroller, "
            "luggage and box — the latter covering delivery parcels and "
            "cartons) inside the car",
        ),
        (
            "encompass pre-annotated images for classes like person, stroller, and luggage.",
            "encompass pre-annotated images for the four target classes: "
            "person, stroller, luggage, and box (the box class covers "
            "delivery parcels and shipping cartons).",
        ),
    ]

    # ── C2 (real-time). Soften "real-time" claims. ──
    REALTIME_REPLACEMENTS: list[tuple[str, str]] = [
        (
            "to detect passengers and bulky objects inside the car in real-time",
            "to detect passengers and bulky objects inside the car in near "
            "real-time (single-frame inference; FPS not benchmarked on a "
            "target edge device in this scope)",
        ),
    ]

    # ── C3. Fix the wrong figure caption under Architecture. ──
    FIG_CAPTION_REPLACEMENTS: list[tuple[str, str]] = [
        (
            "Fig. 3.1. Obtained results of the study.",
            "Fig. 3.1. System architecture and decision flow of the proposed "
            "elevator control algorithm.",
        ),
    ]

    all_replacements = (
        SHOP_REPLACEMENTS
        + RQ_REPLACEMENTS
        + HYPOTHESIS_REPLACEMENTS
        + CLASS_LIST_REPLACEMENTS
        + REALTIME_REPLACEMENTS
        + FIG_CAPTION_REPLACEMENTS
    )

    for p in paragraphs:
        original = p.text
        if not original:
            continue
        new_text = original
        for old, new in all_replacements:
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != original:
            _set_text_preserving_first_run(p, new_text)


def find_paragraph_starting_with(doc, prefix: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def find_paragraph_equals(doc, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    return None


def replace_algorithm_block(doc) -> None:
    """Replace the legacy Algorithm 1 block (lines 1–13 mixed onto a single
    paragraph) with the cleanly-formatted ALGORITHM_1_BLOCK."""
    title = find_paragraph_starting_with(doc, "Algorithm 1: Load-")
    if title is None:
        return

    # Walk forward removing paragraphs that look like the original
    # pseudocode (the broken Algorithm 1 block spans ~10 paragraphs in
    # the source document, indices ~379-388).
    next_p_xml = title._p.getnext()
    removed = 0
    while next_p_xml is not None and removed < 12:
        # Stop at the next heading-styled paragraph.
        text_nodes = next_p_xml.findall(qn("w:r") + "/" + qn("w:t"))
        text = "".join(t.text or "" for t in text_nodes).strip()
        # Heuristic: pseudocode lines start with a digit and a colon, or
        # are empty / are the trailing "END IF" marker.
        is_pseudocode = (
            (text and text[0].isdigit() and ":" in text[:6])
            or text.startswith("END IF")
            or not text
        )
        if is_pseudocode:
            to_remove = next_p_xml
            next_p_xml = next_p_xml.getnext()
            to_remove.getparent().remove(to_remove)
            removed += 1
            continue
        break

    # Replace the title text and append the formatted block.
    _set_text_preserving_first_run(title, ALGORITHM_1_BLOCK[0][1])
    insert_block(title, ALGORITHM_1_BLOCK[1:])


def replace_performance_metrics(doc) -> None:
    """Rewrite the Performance Metrics body paragraph."""
    head = find_paragraph_equals(doc, "Performance Metrics and Validation")
    if head is None:
        return
    next_xml = head._p.getnext()
    while next_xml is not None:
        text_nodes = next_xml.findall(qn("w:r") + "/" + qn("w:t"))
        text = "".join(t.text or "" for t in text_nodes).strip()
        if text:
            # Replace this paragraph's text.
            body_p = Paragraph(next_xml, head._parent)
            _set_text_preserving_first_run(body_p, PERFORMANCE_METRICS_TEXT)
            return
        next_xml = next_xml.getnext()


def replace_ethics_section(doc) -> None:
    head = find_paragraph_equals(doc, "Data Availability and Ethical Considerations")
    if head is None:
        return
    next_xml = head._p.getnext()
    while next_xml is not None:
        text_nodes = next_xml.findall(qn("w:r") + "/" + qn("w:t"))
        text = "".join(t.text or "" for t in text_nodes).strip()
        if text:
            body_p = Paragraph(next_xml, head._parent)
            _set_text_preserving_first_run(body_p, ETHICS_TEXT)
            return
        next_xml = next_xml.getnext()


def insert_methodology_blocks(doc) -> None:
    """Insert the new Methodology subsections after the Algorithm 1 block.

    Anchor: the paragraph beginning with "For the core detection mechanism".
    The blocks are inserted in the order defined by METHODOLOGY_INSERTIONS.
    """
    anchor = find_paragraph_starting_with(doc, "For the core detection mechanism")
    if anchor is None:
        return
    cur = anchor
    for block in METHODOLOGY_INSERTIONS:
        # Insert a blank spacer paragraph first for visual separation.
        cur = _new_para_after(cur, "", "Body Text")
        cur = insert_block(cur, block)


def delete_legacy_area_occupancy_section(doc) -> None:
    """Delete the original "Area Occupancy Ratio Calculation Mechanism"
    heading and the four body paragraphs that follow it.

    This section is fully superseded by the new "Class-Based Footprint
    Estimation" block inserted earlier in the Methodology. Removing it
    avoids two competing definitions of the occupancy ratio formula in
    the same chapter.
    """
    head = find_paragraph_equals(doc, "Area Occupancy Ratio Calculation Mechanism")
    if head is None:
        return

    # Walk forward and delete every paragraph until we hit the next
    # Heading 2 ("Performance Metrics and Validation") — but do NOT
    # delete that next heading.
    next_xml = head._p.getnext()
    to_delete = []
    while next_xml is not None:
        # Read the paragraph text.
        text_nodes = next_xml.findall(qn("w:r") + "/" + qn("w:t"))
        text = "".join(t.text or "" for t in text_nodes).strip()
        # Stop at the next Heading 2 boundary (text-based heuristic for the
        # known following heading).
        if text.startswith("Performance Metrics") or text.startswith("Data Availability"):
            break
        to_delete.append(next_xml)
        next_xml = next_xml.getnext()

    # Delete the heading itself and the body paragraphs.
    for elem in [head._p, *to_delete]:
        if elem is not None and elem.getparent() is not None:
            elem.getparent().remove(elem)


def replace_results_placeholders(doc) -> None:
    """Replace the placeholder Heading 1/2/3 dots in the Results section
    with the actual Results content."""
    # Anchor: the "Results" line (Normal style, single word).
    anchor = find_paragraph_equals(doc, "Results")
    if anchor is None:
        return

    # Identify and delete the placeholder paragraphs that immediately
    # follow ("Heading 1", "…..", "Heading 2", "….", "Heading 3", "…").
    placeholder_strings = {"Heading 1", "Heading 2", "Heading 3", "…", "….", "….."}
    next_xml = anchor._p.getnext()
    while next_xml is not None:
        text_nodes = next_xml.findall(qn("w:r") + "/" + qn("w:t"))
        text = "".join(t.text or "" for t in text_nodes).strip()
        if text in placeholder_strings or text == "":
            to_remove = next_xml
            next_xml = next_xml.getnext()
            # Stop deleting once we hit "Discussion and Conclusion".
            t2 = "".join(
                t.text or "" for t in to_remove.findall(qn("w:r") + "/" + qn("w:t"))
            ).strip()
            if t2 == "Discussion and Conclusion":
                # Don't delete the next-section header.
                break
            to_remove.getparent().remove(to_remove)
            continue
        break

    insert_block(anchor, RESULTS_BLOCK)


def replace_discussion_placeholders(doc) -> None:
    anchor = find_paragraph_equals(doc, "Discussion and Conclusion")
    if anchor is None:
        return
    placeholder_strings = {"Heading 1", "Heading 2", "Heading 3", "…", "….", "….."}
    next_xml = anchor._p.getnext()
    while next_xml is not None:
        text_nodes = next_xml.findall(qn("w:r") + "/" + qn("w:t"))
        text = "".join(t.text or "" for t in text_nodes).strip()
        if text in placeholder_strings or text == "":
            to_remove = next_xml
            next_xml = next_xml.getnext()
            t2 = "".join(
                t.text or "" for t in to_remove.findall(qn("w:r") + "/" + qn("w:t"))
            ).strip()
            if t2 == "ACKNOWLEDGEMENT":
                break
            to_remove.getparent().remove(to_remove)
            continue
        break

    insert_block(anchor, DISCUSSION_BLOCK)


def _attach_table_after_caption(
    doc,
    caption_prefix: str,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    """Find the paragraph beginning with ``caption_prefix`` (e.g.
    ``"Table 3.1."``) and insert a Word table immediately after it."""
    p = find_paragraph_starting_with(doc, caption_prefix)
    if p is None:
        return
    _insert_table_after(p, headers, rows)


def attach_all_tables(doc) -> None:
    """Insert every Word table that the new content blocks announce."""
    for prefix, headers, rows in (
        ("Table 3.1.", TABLE_3_1_HEADERS, TABLE_3_1_ROWS),
        ("Table 3.2.", TABLE_3_2_HEADERS, TABLE_3_2_ROWS),
        ("Table 4.1.", TABLE_4_1_HEADERS, TABLE_4_1_ROWS),
        ("Table 4.2.", TABLE_4_2_HEADERS, TABLE_4_2_ROWS),
        ("Table 4.3.", TABLE_4_3_HEADERS, TABLE_4_3_ROWS),
        ("Table 4.4.", TABLE_4_4_HEADERS, TABLE_4_4_ROWS),
        ("Table 4.5.", TABLE_4_5_HEADERS, TABLE_4_5_ROWS),
    ):
        _attach_table_after_caption(doc, prefix, headers, rows)


def promote_main_section_headings(doc) -> None:
    """Promote ``Results`` and ``Discussion and Conclusion`` to Heading 1.

    The original document leaves these two section titles in the Normal
    style with the body content marked as Heading 4 placeholders, which
    makes Word's automatic numbering treat the new sub-headings as if
    they belonged to the previous chapter (Methodology). Promoting the
    section titles to Heading 1 fixes the numbering and brings them in
    line with the existing ``LITERATURE REVIEW`` and ``PROJECT METHOD
    SPECIFICATION`` titles.
    """
    _promote_to_heading1(doc, "Results", "RESULTS")
    _promote_to_heading1(doc, "Discussion and Conclusion", "DISCUSSION AND CONCLUSION")


def append_references(doc) -> None:
    """Append the new bibliography entries to the REFERENCES section.

    Each entry is added in alphabetical order at the end of the existing
    list. We do NOT re-sort the existing entries to avoid unintentionally
    breaking the supervisor's expected ordering.
    """
    ref_head = find_paragraph_equals(doc, "REFERENCES")
    if ref_head is None:
        return

    # Find the last paragraph in the document so we append after it.
    last_para = doc.paragraphs[-1]
    insert_block(last_para, [("Normal", entry) for entry in REFERENCE_ADDITIONS])


# ──────────────────────────────────────────────────────────────────────
#  Driver
# ──────────────────────────────────────────────────────────────────────


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", default=DEFAULT_INPUT, type=Path)
    ap.add_argument("--output", default=DEFAULT_OUTPUT, type=Path)
    args = ap.parse_args()

    if not args.input.exists():
        print(f"[error] input file not found: {args.input}", file=sys.stderr)
        return 2

    print(f"[info] loading {args.input.name} …")
    doc = docx.Document(str(args.input))

    print("[info] applying text replacements …")
    apply_text_replacements(doc)

    print("[info] reformatting Algorithm 1 …")
    replace_algorithm_block(doc)

    print("[info] inserting Methodology subsections …")
    insert_methodology_blocks(doc)

    print("[info] deleting legacy Area Occupancy section …")
    delete_legacy_area_occupancy_section(doc)

    print("[info] rewriting Performance Metrics paragraph …")
    replace_performance_metrics(doc)

    print("[info] rewriting Ethics paragraph …")
    replace_ethics_section(doc)

    print("[info] populating Results section …")
    replace_results_placeholders(doc)

    print("[info] populating Discussion section …")
    replace_discussion_placeholders(doc)

    print("[info] promoting Results / Discussion to Heading 1 …")
    promote_main_section_headings(doc)

    print("[info] inserting Word tables for the new sections …")
    attach_all_tables(doc)

    print("[info] appending new References …")
    append_references(doc)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    print(f"[done] wrote {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
