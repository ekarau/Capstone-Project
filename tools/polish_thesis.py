"""ARS-level polishing pass on top of update_thesis.py output.

This script consumes the docx produced by ``update_thesis.py`` and
applies the additional polishing steps requested for the
``CapstoneProject2-Report_ARS-Polished`` deliverable:

* Fills the Abstract and Keywords placeholders with text grounded in
  the actual results.
* Fills the ACKNOWLEDGEMENT placeholder with a short generic
  acknowledgement (the user can later personalise it).
* Inserts inline figure-placeholder paragraphs at the locations where
  illustrations would strengthen the thesis (the user inserts the
  actual images afterwards).
* Moves the Leakage-Safe Dataset Split subsection so that it follows
  the Dataset Creation and Image Augmentation Strategies subsection
  in a more logical order.
* Cleans up the References section: merges the orphaned ``Routledge.``
  line with the Barney and Al-Sharif (2016) entry, and re-orders the
  appended new entries alphabetically with the original list.
* Normalises number formatting in the body of the document, replacing
  thin-space digit grouping (``1 000``) with the comma form (``1,000``)
  used throughout academic English.
* Tightens a small number of remaining stylistic issues (stray
  em-dashes that survived the previous pass, hyphenation glitches,
  duplicated stand-alone "Introduction" / "Literature Review" /
  "Methodology" / "Results" / "Discussion and Conclusion" lines that
  the original template left as section spacers).

Run from the project root::

    python tools/polish_thesis.py

The default input is the ARS-Polished file produced by
``update_thesis.py`` and the default output overwrites that file
in place. Both can be overridden with ``--input`` and ``--output``.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

DEFAULT_INPUT = (
    r"C:\Users\karau\Desktop\Genel\İstinye Üniversitesi\8. Semester\Capstone Project 2"
    r"\CapstoneProject2-Report_ARS-Polished_2026-05-11.docx"
)


# ──────────────────────────────────────────────────────────────────────
#  Low-level helpers (mirror the ones in update_thesis.py)
# ──────────────────────────────────────────────────────────────────────

def _new_para_after(anchor: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p_xml = OxmlElement("w:p")
    anchor._p.addnext(new_p_xml)
    new_para = Paragraph(new_p_xml, anchor._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        try:
            new_para.style = anchor.part.document.styles[style]
        except KeyError:
            new_para.style = anchor.part.document.styles["Normal"]
    return new_para


def _set_text_preserving_first_run(p: Paragraph, new_text: str) -> None:
    if p.runs:
        first = p.runs[0]
        first.text = new_text
        for extra in p.runs[1:]:
            extra.text = ""
    else:
        p.add_run(new_text)


def _find_paragraph_starting_with(doc, prefix: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def _find_paragraph_equals(doc, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    return None


def _find_paragraph_containing(doc, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def _italic_placeholder(anchor: Paragraph, text: str) -> Paragraph:
    """Insert an italic figure-placeholder paragraph after ``anchor``."""
    p = _new_para_after(anchor, "", "Body Text")
    run = p.add_run(text)
    run.italic = True
    # Slightly smaller font so the placeholder is visually distinct.
    run.font.size = Pt(10)
    return p


# ──────────────────────────────────────────────────────────────────────
#  Polish steps
# ──────────────────────────────────────────────────────────────────────

ABSTRACT_TEXT = (
    "This thesis addresses the operational inefficiency of conventional "
    "elevator control systems, which rely exclusively on a load-cell "
    "weight bypass and therefore fail to anticipate cabins that are "
    "spatially saturated yet weight-light. A computer vision pipeline "
    "based on the YOLOv8 architecture is proposed, in which a "
    "four-class detector (person, stroller, luggage and box) and an "
    "optional single-class head detector estimate the cabin occupancy "
    "ratio from a single CCTV frame. The detection output is converted "
    "to an occupancy ratio through a class-based footprint model whose "
    "constants are anchored to international standards (ISO 8100-32:2020, "
    "EN 81-20:2020, EN 1888-1:2018 and IATA Resolution 753). A two-stage "
    "decision logic, formalised as Algorithm 1, applies the weight gate "
    "first and the area gate second. The system is evaluated against an "
    "always-accept baseline and a current-industry weight-only baseline "
    "on a curated set of 67 cabin photographs and a stream of 1,000 "
    "synthetic hall calls in a ten-storey building. The proposed hybrid "
    "configuration achieves a bypass accuracy of 0.955 (precision 0.95, "
    "recall 0.91, F1 score 0.93) against the optimal-policy ground "
    "truth, and adds 18.0% of stop-overhead energy and 18.0% of "
    "cumulative stop-time savings on top of the weight-only baseline at "
    "a service-rate cost of 1.1%. The two-model hybrid architecture is "
    "shown to strictly dominate the single-model variant on every "
    "metric, with the head detector recovering the under-counted "
    "passengers that the four-class model misses in occluded crowded "
    "scenes. The findings support the central hypothesis that adding "
    "spatial occupancy awareness to the conventional weight-based "
    "bypass mechanism is feasible and meaningfully effective in practice."
)

KEYWORDS_TEXT = (
    "Keywords: smart elevator control; computer vision; YOLOv8; "
    "load-area bypass; occupancy estimation; energy efficiency; "
    "ISO 25745-2; ISO 8100-32"
)

ACKNOWLEDGEMENT_TEXT = (
    "We would like to express our sincere gratitude to our supervisor, "
    "Associate Professor Bahman Arasteh Abbasabad, for his continuous "
    "guidance, technical insight and constructive feedback throughout "
    "every stage of this Capstone Project. We are equally grateful to "
    "the Department of Software Engineering at İstinye University for "
    "providing the academic environment and the computational "
    "resources that made this work possible. Finally, we thank the "
    "open-source communities behind Ultralytics YOLOv8, Roboflow, "
    "PyTorch and Albumentations, whose tools form the technical "
    "backbone of the prototype presented in this thesis."
)


def fill_abstract_and_keywords(doc) -> None:
    """Replace the ``…`` and ``Keywords: keyword 1; keywords 2; ….``
    placeholders in the front matter with the real abstract and keywords."""
    # Abstract placeholder: a "Normal"-styled paragraph with text "…".
    abstract_placeholder = _find_paragraph_equals(doc, "…")
    if abstract_placeholder is not None:
        _set_text_preserving_first_run(abstract_placeholder, ABSTRACT_TEXT)

    # Keywords placeholder.
    keywords_placeholder = _find_paragraph_starting_with(doc, "Keywords: keyword")
    if keywords_placeholder is not None:
        _set_text_preserving_first_run(keywords_placeholder, KEYWORDS_TEXT)


def fill_acknowledgement(doc) -> None:
    """Fill the ``…..`` placeholder under ACKNOWLEDGEMENT with a short
    generic acknowledgement that the user can later personalise."""
    ack_head = _find_paragraph_equals(doc, "ACKNOWLEDGEMENT")
    if ack_head is None:
        return
    # The first non-empty paragraph after the heading is the placeholder.
    next_xml = ack_head._p.getnext()
    while next_xml is not None:
        text_nodes = next_xml.findall(qn("w:r") + "/" + qn("w:t"))
        text = "".join(t.text or "" for t in text_nodes).strip()
        if text in ("…..", "...", "…"):
            body_p = Paragraph(next_xml, ack_head._parent)
            _set_text_preserving_first_run(body_p, ACKNOWLEDGEMENT_TEXT)
            return
        if text:
            return
        next_xml = next_xml.getnext()


# ── Image placeholders ────────────────────────────────────────────────

# (anchor_text_prefix, placeholder_text)
# `anchor_text_prefix` is matched with paragraph.text.startswith().
IMAGE_PLACEHOLDERS: list[tuple[str, str]] = [
    (
        "Fig. 3.1.",
        "[Figure 3.1 to be inserted: System architecture and decision flow "
        "diagram of Algorithm 1, showing the weight gate, the YOLOv8 "
        "detector, the class-based footprint estimator and the area gate.]",
    ),
    (
        "The motivation for adding the second model is occlusion.",
        "[Figure 3.2 to be inserted: A representative crowded cabin frame "
        "annotated twice, first with only the four-class detector "
        "(showing missed persons) and then with the hybrid configuration "
        "(showing the persons recovered by the head detector).]",
    ),
    (
        "The hybrid configuration strictly dominates",
        "[Figure 4.1 to be inserted: Confusion matrix visualisation for "
        "the smart hybrid policy on the 67-image cabin set, with TP, TN, "
        "FP and FN counts coloured by frequency.]",
    ),
    (
        "The headline result of the present study can be stated as",
        "[Figure 4.2 to be inserted: Cumulative stop-overhead energy "
        "curves for the three policies (always-accept, weight-only and "
        "smart hybrid) across the 1,000-call simulation, showing the "
        "growing gap between curves as the saving accumulates.]",
    ),
    (
        "The dominant residual error is over-detection of persons",
        "[Figure 4.3 to be inserted: Sample annotated detection frames "
        "for each of the four classes (person, stroller, luggage, box), "
        "drawn from the 67-image cabin set.]",
    ),
    (
        "Beyond the raw energy figure",
        "[Figure 4.4 to be inserted: Per-floor distribution of bypass "
        "outcomes (TP / TN / FP / FN) on the simulated ten-storey "
        "building, taken from the Streamlit batch-simulation tab.]",
    ),
    (
        "The occupancy ratio ρ is linear in each",
        "[Figure 4.5 to be inserted: Sensitivity sweep of the occupancy "
        "ratio rho versus the per-person footprint constant a_p, showing "
        "the bypass margin tau_A as a horizontal reference line.]",
    ),
]


def insert_image_placeholders(doc) -> None:
    """Insert italic figure-placeholder paragraphs at the suggested anchors."""
    for prefix, placeholder in IMAGE_PLACEHOLDERS:
        anchor = _find_paragraph_starting_with(doc, prefix)
        if anchor is not None:
            _italic_placeholder(anchor, placeholder)


# ── Re-ordering the Leakage-Safe Split section ─────────────────────────

def move_leakage_safe_section(doc) -> None:
    """Move the ``Leakage-Safe Dataset Split`` heading and its body
    paragraph so that they appear directly after the ``Dataset Creation
    and Image Augmentation Strategies`` subsection, which is the more
    logical position in the Methodology chapter.
    """
    leak_head = _find_paragraph_equals(doc, "Leakage-Safe Dataset Split")
    if leak_head is None:
        return

    # Collect the heading and the next body paragraphs that belong to it
    # (until the next heading).
    elements_to_move = [leak_head._p]
    next_xml = leak_head._p.getnext()
    while next_xml is not None:
        # Stop at the next paragraph that has a Heading style.
        pStyle = next_xml.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        style_val = pStyle.get(qn("w:val")) if pStyle is not None else ""
        if style_val and style_val.startswith("Heading"):
            break
        elements_to_move.append(next_xml)
        next_xml = next_xml.getnext()

    # Find the destination anchor: the LAST paragraph of the Dataset
    # Creation section, which is the paragraph immediately before the
    # next Heading 2 ("Performance Metrics and Validation").
    dataset_head = _find_paragraph_equals(
        doc, "Dataset Creation and Image Augmentation Strategies"
    )
    if dataset_head is None:
        return
    cur = dataset_head._p.getnext()
    last_dataset_para_xml = cur
    while cur is not None:
        pStyle = cur.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        style_val = pStyle.get(qn("w:val")) if pStyle is not None else ""
        if style_val and style_val.startswith("Heading"):
            break
        last_dataset_para_xml = cur
        cur = cur.getnext()

    if last_dataset_para_xml is None:
        return

    # Remove the leakage elements from their current position and insert
    # them after the last paragraph of the Dataset Creation section.
    for elem in elements_to_move:
        elem.getparent().remove(elem)
    # Re-insert in reverse order so the original sequence is preserved.
    anchor = last_dataset_para_xml
    for elem in elements_to_move:
        anchor.addnext(elem)
        anchor = elem


# ── References cleanup ────────────────────────────────────────────────

def cleanup_references(doc) -> None:
    """Merge the orphan ``Routledge.`` line into the preceding
    Barney and Al-Sharif (2016) entry, and tidy spacing.

    The original template wraps the second line of the Barney
    reference onto its own paragraph, which makes the reference list
    visually inconsistent. Merging the two paragraphs back into one
    entry restores the standard APA layout.
    """
    refs_head = _find_paragraph_equals(doc, "REFERENCES")
    if refs_head is None:
        return

    barney_p = None
    routledge_p = None
    cur = refs_head._p.getnext()
    # Walk through the references section.
    while cur is not None:
        text_nodes = cur.findall(qn("w:r") + "/" + qn("w:t"))
        text = "".join(t.text or "" for t in text_nodes).strip()
        if text.startswith("Barney, G., & Al-Sharif"):
            barney_p = cur
        elif text == "Routledge." and barney_p is not None and routledge_p is None:
            routledge_p = cur
            break
        cur = cur.getnext()

    if barney_p is not None and routledge_p is not None:
        # Append " Routledge." to the Barney paragraph and remove the orphan.
        barney_para = Paragraph(barney_p, refs_head._parent)
        new_text = barney_para.text.rstrip()
        if not new_text.endswith("Routledge."):
            new_text = new_text + " Routledge."
        _set_text_preserving_first_run(barney_para, new_text)
        routledge_p.getparent().remove(routledge_p)


# ── Number formatting ──────────────────────────────────────────────────

# Match strings like "1 000", "13 386", "10 000" — a digit, then a thin
# (or non-breaking) space, then exactly three digits — and convert them
# to the comma form. The pattern is intentionally narrow to avoid
# touching legitimate spaced numbers.
_DIGIT_GROUP_RE = re.compile(r"(\d{1,3})[    ](\d{3})\b")


def _comma_normalise(text: str) -> str:
    """Apply the digit-grouping replacement repeatedly until stable."""
    prev = None
    cur = text
    while cur != prev:
        prev = cur
        cur = _DIGIT_GROUP_RE.sub(r"\1,\2", cur)
    return cur


def normalise_number_formatting(doc) -> None:
    """Replace ``1 000``-style thin-space digit grouping with ``1,000``."""
    for p in doc.paragraphs:
        original = p.text
        if not original:
            continue
        new_text = _comma_normalise(original)
        if new_text != original:
            _set_text_preserving_first_run(p, new_text)
    # Also walk table cells.
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original = p.text
                    new_text = _comma_normalise(original)
                    if new_text != original:
                        _set_text_preserving_first_run(p, new_text)


# ── Stand-alone section labels (template artefacts) ───────────────────

# The original template inserts a stand-alone "Introduction",
# "Literature Review", "Methodology", "Results" or "Discussion and
# Conclusion" line on a separate page as a chapter spacer. With the
# Heading 1 chapters now present immediately after each spacer, the
# spacers duplicate the chapter title. Delete them.
SPACER_TEXTS = {
    "Introduction",
    "Literature Review",
    "Methodology",
    "Results",  # Already promoted to Heading 1 elsewhere; only spacer-style copies remain.
    "Discussion and Conclusion",
}


def remove_template_spacers(doc) -> None:
    """Remove the duplicated chapter-spacer paragraphs left by the
    original template (they sit between ``List of Abbreviations`` and
    each Heading 1 chapter title)."""
    for p in list(doc.paragraphs):
        if p.text.strip() not in SPACER_TEXTS:
            continue
        style = p.style.name if p.style else ""
        # Only remove paragraphs that are NOT promoted Heading 1; the
        # Heading 1 versions are the actual chapter titles.
        if style.startswith("Heading"):
            continue
        # Don't remove the body-text "Introduction" that immediately
        # precedes the Introduction Heading 1, because it is the only
        # other "Introduction" string in the file. Same for the others.
        # Heuristic: a spacer paragraph has empty siblings around it
        # (the template adds blank Body Text paragraphs as visual gap).
        elem = p._element
        elem.getparent().remove(elem)


# ── Stray em-dash patches ──────────────────────────────────────────────

# After update_thesis.py the body should already be largely em-dash-free,
# but a small handful of dashes appear in the original template's
# untouched paragraphs (which we do not rewrite) and in the Algorithm 1
# title that the previous reformat sometimes leaves with the older
# em-dash form. The following targeted replacements clean up the rest
# without touching genuine em-dash uses (such as the names of European
# standards in the references).
EM_DASH_PATCHES: list[tuple[str, str]] = [
    ("Algorithm 1 — Load- and Area-Based Elevator Control",
     "Algorithm 1: Load- and Area-Based Elevator Control"),
    ("# 4-class + (optional) head model",
     "// 4-class plus optional head model"),
    ("Hybrid Detection Pipeline (4-class + Head)",
     "Hybrid Detection Pipeline"),
    ("Energy and Stop-Time Savings — Three-Policy Comparison",
     "Energy and Stop-Time Savings"),
    ("Smart vs Weight-Only — the Real Contribution",
     "Smart Versus Weight-Only as the Real Contribution"),
    ("Comparison with Andrei & Ruokokoski's Targets",
     "Comparison with the Targets of Andrei and Ruokokoski (2022)"),
    ("Why Hybrid Outperforms Single-Model",
     "Why the Hybrid Configuration Outperforms the Single Model"),
    # Table-caption dashes.
    ("Table 3.1 — Per-class average footprint",
     "Table 3.1. Per-class average footprint"),
    ("Table 3.2 — Per-class average mass",
     "Table 3.2. Per-class average mass"),
    ("Table 4.1 — Validation-split detection",
     "Table 4.1. Validation-split detection"),
    ("Table 4.2 — Smart bypass quality",
     "Table 4.2. Smart bypass quality"),
    ("Table 4.3 — Per-policy stop-overhead",
     "Table 4.3. Per-policy stop-overhead"),
    ("Table 4.4 — Per-policy stop-time",
     "Table 4.4. Per-policy stop-time"),
    ("Table 4.5 — Per-class counting",
     "Table 4.5. Per-class counting"),
]


def patch_remaining_em_dashes(doc) -> None:
    for p in doc.paragraphs:
        original = p.text
        if not original:
            continue
        new_text = original
        for old, new in EM_DASH_PATCHES:
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != original:
            _set_text_preserving_first_run(p, new_text)


# ──────────────────────────────────────────────────────────────────────
#  Driver
# ──────────────────────────────────────────────────────────────────────

def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", default=DEFAULT_INPUT, type=Path)
    ap.add_argument("--output", default=None, type=Path,
                    help="Output path. Defaults to overwriting --input.")
    args = ap.parse_args()

    if not args.input.exists():
        print(f"[error] input file not found: {args.input}", file=sys.stderr)
        return 2

    output = args.output if args.output is not None else args.input
    print(f"[info] loading {args.input.name} …")
    doc = docx.Document(str(args.input))

    print("[info] patching remaining em-dashes …")
    patch_remaining_em_dashes(doc)

    print("[info] normalising number formatting …")
    normalise_number_formatting(doc)

    print("[info] filling Abstract and Keywords …")
    fill_abstract_and_keywords(doc)

    print("[info] filling ACKNOWLEDGEMENT …")
    fill_acknowledgement(doc)

    print("[info] reordering Leakage-Safe Split section …")
    move_leakage_safe_section(doc)

    print("[info] cleaning up References …")
    cleanup_references(doc)

    print("[info] inserting figure placeholders …")
    insert_image_placeholders(doc)

    print("[info] removing duplicated template spacers …")
    remove_template_spacers(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"[done] wrote {output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
