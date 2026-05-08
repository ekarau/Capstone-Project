"""Generate a Word (.docx) report of the unified dataset class distribution.

Output: results/dataset_class_distribution.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


HEADER_FILL = "1F4E78"   # dark blue
ZEBRA_FILL = "F2F2F2"    # light grey for alternating rows
TOTAL_FILL = "DCE6F1"    # light blue for total rows


def set_cell_fill(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    color_rgb: RGBColor | None = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    font_size: int = 10,
    fill: str | None = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.name = "Calibri"
    if color_rgb is not None:
        run.font.color.rgb = color_rgb
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        set_cell_fill(cell, fill)
    set_cell_borders(cell)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 0:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def make_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    col_widths_cm: list[float] | None = None,
    total_row_indices: tuple[int, ...] = (),
    align_right_cols: tuple[int, ...] = (),
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False

    # Header row
    for col_idx, header in enumerate(headers):
        write_cell(
            table.cell(0, col_idx),
            header,
            bold=True,
            color_rgb=RGBColor(0xFF, 0xFF, 0xFF),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10,
            fill=HEADER_FILL,
        )

    # Body rows
    for row_idx, row in enumerate(rows):
        is_total = row_idx in total_row_indices
        fill = TOTAL_FILL if is_total else (ZEBRA_FILL if row_idx % 2 == 1 else None)
        for col_idx, value in enumerate(row):
            align = (
                WD_ALIGN_PARAGRAPH.RIGHT
                if col_idx in align_right_cols
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            write_cell(
                table.cell(row_idx + 1, col_idx),
                value,
                bold=is_total,
                align=align,
                font_size=10,
                fill=fill,
            )

    # Column widths
    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(w)

    doc.add_paragraph()  # spacing after table


def main() -> None:
    out_dir = Path(__file__).resolve().parent.parent / "results"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "dataset_class_distribution.docx"

    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("Final Unified Dataset — Class Distribution", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    add_paragraph(
        doc,
        "Smart Elevator CV — Capstone Project | Computer-vision-based occupancy "
        "estimation for elevator control. The tables below summarize the unified "
        "training dataset after class remap, leakage-safe group-stratified split, "
        "and offline class-balancing augmentation.",
        italic=True,
    )

    # ─────────────────────────────────────────────────────────
    # Table 1 — Overall (after augmentation)
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "1. Overall Class Distribution (after augmentation)", level=1)

    headers_1 = ["Class", "Train", "Val", "Test", "TOTAL Instances"]
    rows_1 = [
        ["person",   "20,591", "1,215", "844", "22,650"],
        ["stroller",  "8,073",   "587", "178",  "8,838"],
        ["luggage",   "1,757",   "204",  "62",  "2,023"],
        ["box",       "1,432",   "121",  "40",  "1,593"],
        ["TOTAL Instances", "31,853", "2,127", "1,124", "35,104"],
        ["Image count",     "13,386",   "857",   "367", "14,610"],
    ]
    make_table(
        doc,
        headers_1,
        rows_1,
        col_widths_cm=[3.5, 2.6, 2.0, 2.0, 3.5],
        total_row_indices=(4, 5),
        align_right_cols=(1, 2, 3, 4),
    )

    # ─────────────────────────────────────────────────────────
    # Table 2 — Augmentation Multiplier Effect
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "2. Train Detail — Before vs After Augmentation", level=1)

    headers_2 = [
        "Class",
        "Train (original)",
        "Aug Multiplier",
        "Train (after aug)",
        "Increase",
    ]
    rows_2 = [
        ["person",   "6,931",  "3x",  "20,591", "+197%"],
        ["stroller", "2,989",  "2x",   "8,073", "+170%"],
        ["luggage",    "829",  "2x",   "1,757", "+112%"],
        ["box",        "287",  "5x",   "1,432", "+399%"],
    ]
    make_table(
        doc,
        headers_2,
        rows_2,
        col_widths_cm=[3.0, 3.0, 3.0, 3.5, 2.5],
        align_right_cols=(1, 2, 3, 4),
    )

    # ─────────────────────────────────────────────────────────
    # Table 3 — Val and Test (no augmentation)
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "3. Val and Test — Clean Splits (no augmentation)", level=1)

    headers_3 = ["Class", "Val", "Test"]
    rows_3 = [
        ["person",   "1,215", "844"],
        ["stroller",   "587", "178"],
        ["luggage",    "204",  "62"],
        ["box",        "121",  "40"],
    ]
    make_table(
        doc,
        headers_3,
        rows_3,
        col_widths_cm=[5.0, 3.5, 3.5],
        align_right_cols=(1, 2),
    )

    # ─────────────────────────────────────────────────────────
    # Table 4 — Per-Source Image Distribution
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "4. Per-Source Image Distribution", level=1)

    headers_4 = ["Source", "Train", "Val", "Test", "Notes"]
    rows_4 = [
        ["Elevator.yolov8",                "0",     "0",   "80",  "Custom elevator data (test-only)"],
        ["-People Counting.yolov8",        "514",  "96",   "32",  "Person video frames"],
        ["people ditection in elevator",   "197",   "0",    "0",  "Small set, all to train"],
        ["top down view.yolov8",            "39",   "0",    "0",  "Very small, all to train"],
        ["normal3.yolov8",                  "85",  "16",    "5",  "908 empty labels dropped"],
        ["Stroller.yolov8",              "2,676", "501",  "168",  "Multi-class source"],
        ["My Luggage.yolov8",               "66",  "13",    "1",  "—"],
        ["luggage.yolov8",                  "80",  "15",    "5",  "Polygon→bbox converted"],
        ["suitcase.yolov8",                "446",  "83",   "28",  "—"],
        ["box.yolov8",                     "121",  "23",   "11",  "Augmentation 5x applied"],
        ["lastdataset_extra.yolov8",       "590", "110",   "37",  "★ Newly added (deduplicated)"],
        ["TOTAL (pre-aug)",              "4,814", "857",  "367",  "—"],
    ]
    make_table(
        doc,
        headers_4,
        rows_4,
        col_widths_cm=[5.0, 1.7, 1.5, 1.5, 6.5],
        total_row_indices=(11,),
        align_right_cols=(1, 2, 3),
    )

    # ─────────────────────────────────────────────────────────
    # Table 5 — Quality Checks
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "5. Data Quality Checks", level=1)

    headers_5 = ["Check", "Result"]
    rows_5 = [
        ["Empty labels (dropped)",                      "957"],
        ["Class-map drops (dropped)",                   "2"],
        ["Group-prefix leakage (between splits)",       "0   ✅ none"],
        ["Augmented files in Val",                      "0   ✅ clean"],
        ["Augmented files in Test",                     "0   ✅ clean"],
        ["Augmented files in Train",                    "8,572"],
    ]
    make_table(
        doc,
        headers_5,
        rows_5,
        col_widths_cm=[10.0, 6.0],
    )

    # ─────────────────────────────────────────────────────────
    # Table 6 — Bonus (Thesis §3.10 cross-reference)
    # ─────────────────────────────────────────────────────────
    add_heading(doc, "6. Cross-Reference with Thesis §3.10", level=1)

    headers_6 = ["Class", "Spatial Unit (Thesis)", "Aug Multiplier", "Test Instances"]
    rows_6 = [
        ["person",   "20,000", "3x",  "844"],
        ["stroller", "45,000", "2x",  "178"],
        ["luggage",  "30,000", "2x",   "62"],
        ["box",      "25,000", "5x",   "40"],
    ]
    make_table(
        doc,
        headers_6,
        rows_6,
        col_widths_cm=[3.5, 4.5, 4.0, 4.0],
        align_right_cols=(1, 2, 3),
    )

    # Footer note
    add_heading(doc, "Notes", level=2)
    add_paragraph(
        doc,
        "• Group-stratified split: video frames sharing a stem prefix are placed "
        "in a single split to prevent near-duplicate leakage between train, val, "
        "and test.",
    )
    add_paragraph(
        doc,
        "• Augmentation operates ONLY on the train split; val and test are kept "
        "fully clean to ensure metrics reflect true generalization.",
    )
    add_paragraph(
        doc,
        "• The Elevator.yolov8 source represents the user's deployment scenario "
        "(Manolya Evleri & Cemevi cabin imagery) and is held out 100% as test.",
    )

    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Size : {out_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
