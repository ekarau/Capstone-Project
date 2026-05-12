"""Template- and guideline-compliance pass for the capstone thesis.

This is the third polishing script in the pipeline. It consumes the
ARS-Polished docx produced by ``polish_thesis.py`` and brings the file
into closer conformance with the official Capstone Project Final
Report template and the COE401/COE402 writing guidelines.

Specifically:

* Sets the default body font to Times New Roman 12pt with 1.5 line
  spacing, and sets the page margins to 1 inch on all sides.
* Reduces the keyword list to the maximum of five permitted by the
  guideline.
* Sorts the References section alphabetically (case-insensitive) by
  first-word author surname, which the previous append-only passes
  could not preserve.
* Inserts a real Word ``TOC`` field after the ``Table of Contents``
  heading. When the user opens the file in Word and right-clicks
  ``Update Field``, the table of contents is populated automatically
  from the Heading 1, 2 and 3 styles already present in the body.
* Inserts a Word ``TOC \\h \\z \\t "Caption,1"`` field after the
  ``List of Tables`` and ``List of Figures`` headings, and applies
  the ``Caption`` style to the existing static table and figure
  captions so that the two lists can be auto-populated.
* Adds a hard page break before each Heading 1 chapter title so the
  five chapters always start on a fresh page (consistent with the
  template).
* Writes the result to a new file
  ``CapstoneProject2-Report_Template-Compliant_2026-05-11.docx``
  alongside the input.
"""

from __future__ import annotations

import argparse
import contextlib
import sys
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

DEFAULT_INPUT = (
    r"C:\Users\karau\Desktop\Genel\İstinye Üniversitesi\8. Semester\Capstone Project 2"
    r"\CapstoneProject2-Report_ARS-Polished_2026-05-11.docx"
)
DEFAULT_OUTPUT = (
    r"C:\Users\karau\Desktop\Genel\İstinye Üniversitesi\8. Semester\Capstone Project 2"
    r"\CapstoneProject2-Report_Template-Compliant_2026-05-11.docx"
)


# ──────────────────────────────────────────────────────────────────────
#  Low-level helpers
# ──────────────────────────────────────────────────────────────────────

def _find_paragraph_equals(doc, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    return None


def _find_paragraph_starting_with(doc, prefix: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def _set_text_preserving_first_run(p: Paragraph, new_text: str) -> None:
    if p.runs:
        first = p.runs[0]
        first.text = new_text
        for extra in p.runs[1:]:
            extra.text = ""
    else:
        p.add_run(new_text)


def _ensure_style(doc, style_name: str, base_style: str = "Normal"):
    """Return the named style, creating a passthrough copy of base_style if absent."""
    try:
        return doc.styles[style_name]
    except KeyError:
        pass
    # Create a paragraph style derived from the base style.
    from docx.enum.style import WD_STYLE_TYPE

    new_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    new_style.base_style = doc.styles[base_style]
    return new_style


# ──────────────────────────────────────────────────────────────────────
#  Document defaults: font, line spacing, margins
# ──────────────────────────────────────────────────────────────────────

def set_document_defaults(doc) -> None:
    """Set Times New Roman 12pt with 1.5 line spacing on the Normal
    style, and 1 inch margins on every section.

    Per-paragraph overrides applied later (e.g. abstract, keywords,
    captions, references at 11pt) win against this default.
    """
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    # Apply font name to East-Asian text variant as well so Word does
    # not fall back to Calibri for special characters.
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")

    # 1.5 line spacing on Normal.
    normal.paragraph_format.line_spacing = 1.5

    # Body Text style — used by most of our content paragraphs.
    if "Body Text" in [s.name for s in doc.styles]:
        body = doc.styles["Body Text"]
        body.font.name = "Times New Roman"
        body.font.size = Pt(12)
        body.paragraph_format.line_spacing = 1.5

    # Margins on all sections.
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


# ──────────────────────────────────────────────────────────────────────
#  Per-paragraph 11pt overrides for abstract, keywords, captions, refs
# ──────────────────────────────────────────────────────────────────────

def _apply_run_font(p: Paragraph, *, size_pt: int = 11, bold: bool | None = None) -> None:
    """Force every run in ``p`` to Times New Roman, ``size_pt``,
    optionally bold."""
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt)
        # Patch the East-Asian font slot too.
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:cs"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "Times New Roman")
        if bold is not None:
            run.font.bold = bold


def apply_typography_overrides(doc) -> None:
    """Apply the 11pt overrides specified by the guideline."""
    # Abstract paragraph (the long single paragraph filled by polish_thesis.py).
    abstract = _find_paragraph_starting_with(
        doc, "This thesis addresses the operational"
    )
    if abstract is not None:
        _apply_run_font(abstract, size_pt=11)

    # Keywords paragraph.
    keywords = _find_paragraph_starting_with(doc, "Keywords:")
    if keywords is not None:
        _apply_run_font(keywords, size_pt=11)

    # Table and figure captions: bold, 11pt.
    for p in doc.paragraphs:
        text = p.text.strip()
        if (text.startswith("Table ") and (
            text[len("Table "): len("Table ") + 1].isdigit()
        )) or text.startswith("Fig. ") or text.startswith("Figure "):
            _apply_run_font(p, size_pt=11, bold=True)

    # References list: 11pt (not bold). Find REFERENCES heading and walk
    # subsequent paragraphs to the end of the document.
    refs_head = _find_paragraph_equals(doc, "REFERENCES")
    if refs_head is not None:
        cur = refs_head._p.getnext()
        while cur is not None:
            # Only process paragraph elements; skip tables and other XML.
            if cur.tag != qn("w:p"):
                cur = cur.getnext()
                continue
            p = Paragraph(cur, refs_head._parent)
            if p.text and p.text.strip():
                _apply_run_font(p, size_pt=11)
            cur = cur.getnext()


# ──────────────────────────────────────────────────────────────────────
#  Reduce keywords to the five-keyword maximum
# ──────────────────────────────────────────────────────────────────────

KEYWORDS_FIVE = (
    "Keywords: smart elevator control; computer vision; YOLOv8; "
    "load-area bypass; energy efficiency"
)


def reduce_keywords(doc) -> None:
    p = _find_paragraph_starting_with(doc, "Keywords:")
    if p is not None:
        _set_text_preserving_first_run(p, KEYWORDS_FIVE)


# ──────────────────────────────────────────────────────────────────────
#  Hyphenation cleanup (PDF→Word extraction artefacts)
# ──────────────────────────────────────────────────────────────────────

HYPHEN_FIXES: list[tuple[str, str]] = [
    ("Res- YOLOv5", "Res-YOLOv5"),
    ("Open- Sourced", "Open-Sourced"),
    ("13th IEEE Conference", "13th IEEE Conference"),  # spacing safe-no-op
    ("multi- ", "multi-"),
    ("YOLOv1- v3", "YOLOv1–v3"),
    ("real- time", "real-time"),
    ("load- area", "load-area"),
    ("YOLOv5 stands", "YOLOv5 stands"),  # safety no-op
    ("Hue, Saturation, Value", "Hue, Saturation, Value"),  # no-op
    ("anchor- free", "anchor-free"),
]


def fix_hyphenation_artefacts(doc) -> None:
    for p in doc.paragraphs:
        original = p.text
        if not original:
            continue
        new_text = original
        for old, new in HYPHEN_FIXES:
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != original:
            _set_text_preserving_first_run(p, new_text)


# ──────────────────────────────────────────────────────────────────────
#  Sort the References section alphabetically
# ──────────────────────────────────────────────────────────────────────

def _ref_sort_key(text: str) -> tuple[str, str]:
    """Return a case-insensitive sort key for a reference entry.

    The key is the lowercased first word of the entry. For multi-word
    surnames such as "Red Stag Fulfillment" we keep the full first word
    so that the entry sorts under "R" rather than "S".
    """
    cleaned = text.strip().lstrip("  ").lower()
    first_word = cleaned.split()[0] if cleaned else ""
    # Strip trailing punctuation to get a clean key.
    while first_word and not first_word[-1].isalnum():
        first_word = first_word[:-1]
    return (first_word, cleaned)


def _looks_like_reference_continuation(text: str) -> bool:
    """Heuristic: a paragraph in the references region is a continuation
    of the previous entry (rather than a new entry) when its first word
    is purely a year, a digit-led page number, a publisher name on its
    own line, or a lowercase fragment."""
    stripped = text.strip()
    if not stripped:
        return True
    first_word = stripped.split()[0]
    # Year-only continuations such as "2017 13th IEEE Conference..."
    if first_word.isdigit() and len(first_word) == 4:
        return True
    # Lowercase first letter ⇒ continuation.
    if first_word[0].islower():
        return True
    # Single-token publisher line ("Routledge.").
    return len(stripped.split()) <= 2 and stripped.endswith(".")


def sort_references_alphabetically(doc) -> None:
    refs_head = _find_paragraph_equals(doc, "REFERENCES")
    if refs_head is None:
        return

    # Collect all reference paragraphs (after REFERENCES heading until end),
    # merging any continuation lines back into the entry that precedes them.
    ref_xmls: list = []
    last_was_entry: bool = False
    cur = refs_head._p.getnext()
    while cur is not None:
        if cur.tag != qn("w:p"):
            cur = cur.getnext()
            continue
        text_nodes = cur.findall(qn("w:r") + "/" + qn("w:t"))
        text = "".join(t.text or "" for t in text_nodes).strip()
        if not text:
            cur = cur.getnext()
            continue
        if last_was_entry and _looks_like_reference_continuation(text):
            # Merge this continuation back into the previous entry.
            prev = ref_xmls[-1]
            prev_para = Paragraph(prev, refs_head._parent)
            prev_text = prev_para.text.rstrip()
            joiner = "" if prev_text.endswith((" ", "-")) else " "
            _set_text_preserving_first_run(prev_para, prev_text + joiner + text)
            # Remove the continuation paragraph.
            nxt = cur.getnext()
            cur.getparent().remove(cur)
            cur = nxt
            continue
        ref_xmls.append(cur)
        last_was_entry = True
        cur = cur.getnext()

    if not ref_xmls:
        return

    # Sort by first-author surname.
    sorted_xmls = sorted(
        ref_xmls,
        key=lambda elem: _ref_sort_key(
            "".join(
                t.text or ""
                for t in elem.findall(qn("w:r") + "/" + qn("w:t"))
            )
        ),
    )

    # Remove originals, re-insert in sorted order after REFERENCES heading.
    for elem in ref_xmls:
        elem.getparent().remove(elem)

    anchor = refs_head._p
    for elem in sorted_xmls:
        anchor.addnext(elem)
        anchor = elem


# ──────────────────────────────────────────────────────────────────────
#  Word TOC field insertion
# ──────────────────────────────────────────────────────────────────────

def _add_field(paragraph: Paragraph, instr_text: str, placeholder_text: str) -> None:
    """Append a Word field (begin/instruction/separate/placeholder/end)
    to the given paragraph."""
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    placeholder = OxmlElement("w:t")
    placeholder.text = placeholder_text
    run._r.append(placeholder)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _delete_paragraphs_until_heading(start_xml) -> None:
    """Delete paragraphs that follow ``start_xml`` until the next heading."""
    cur = start_xml.getnext()
    while cur is not None:
        # Stop at a Heading-styled paragraph.
        pStyle = cur.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        style_val = pStyle.get(qn("w:val")) if pStyle is not None else ""
        if style_val and style_val.startswith("Heading"):
            return
        # Stop also at a Normal paragraph that contains a chapter / list label.
        text_nodes = cur.findall(qn("w:r") + "/" + qn("w:t"))
        text = "".join(t.text or "" for t in text_nodes).strip()
        if text in (
            "List of Tables",
            "List of Figures",
            "List of Abbreviations",
            "Introduction",
            "Literature Review",
            "Methodology",
            "Results",
            "Discussion and Conclusion",
        ):
            return
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt


def insert_toc_fields(doc) -> None:
    """Replace the empty ``Table of Contents`` / ``List of Tables`` /
    ``List of Figures`` regions with proper Word fields.

    The user must open the produced file in Word, right-click on each
    field and select "Update Field" to populate the lists.
    """
    # ── Table of Contents ──
    toc_head = _find_paragraph_equals(doc, "Table of Contents")
    if toc_head is not None:
        _delete_paragraphs_until_heading(toc_head._p)
        # Insert the field as a new paragraph immediately after the heading.
        new_p_xml = OxmlElement("w:p")
        toc_head._p.addnext(new_p_xml)
        new_para = Paragraph(new_p_xml, toc_head._parent)
        _add_field(
            new_para,
            'TOC \\o "1-3" \\h \\z \\u',
            "Right-click on this line in Microsoft Word and choose "
            '"Update Field" to populate the table of contents.',
        )

    # ── List of Tables ──
    lot_head = _find_paragraph_equals(doc, "List of Tables")
    if lot_head is not None:
        _delete_paragraphs_until_heading(lot_head._p)
        new_p_xml = OxmlElement("w:p")
        lot_head._p.addnext(new_p_xml)
        new_para = Paragraph(new_p_xml, lot_head._parent)
        _add_field(
            new_para,
            'TOC \\h \\z \\t "Caption,1"',
            "Right-click on this line in Microsoft Word and choose "
            '"Update Field" to populate the list of tables.',
        )

    # ── List of Figures ──
    lof_head = _find_paragraph_equals(doc, "List of Figures")
    if lof_head is not None:
        _delete_paragraphs_until_heading(lof_head._p)
        new_p_xml = OxmlElement("w:p")
        lof_head._p.addnext(new_p_xml)
        new_para = Paragraph(new_p_xml, lof_head._parent)
        _add_field(
            new_para,
            'TOC \\h \\z \\t "Caption,1"',
            "Right-click on this line in Microsoft Word and choose "
            '"Update Field" to populate the list of figures. Note that '
            "the same field is used for tables and figures, since both "
            "are styled with the Caption style; if separate lists are "
            "preferred, switch to label-based filtering with TOC \\c.",
        )


# ──────────────────────────────────────────────────────────────────────
#  Apply Caption style to existing static table/figure captions
# ──────────────────────────────────────────────────────────────────────

def apply_caption_style(doc) -> None:
    """Apply the ``Caption`` style to every paragraph that begins with
    ``Table N`` or ``Fig. N`` / ``Figure N``."""
    caption_style = _ensure_style(doc, "Caption", base_style="Normal")
    for p in doc.paragraphs:
        text = p.text.strip()
        is_table = (
            text.startswith("Table ")
            and len(text) > len("Table ")
            and text[len("Table ")].isdigit()
        )
        is_figure = (
            text.startswith("Fig. ")
            or text.startswith("Figure ")
            or text.startswith("[Figure ")
        )
        if is_table or is_figure:
            with contextlib.suppress(Exception):
                p.style = caption_style


# ──────────────────────────────────────────────────────────────────────
#  Page breaks before each Heading 1 chapter
# ──────────────────────────────────────────────────────────────────────

CHAPTER_TITLES = {
    "Introduction",
    "LITERATURE REVIEW",
    "PROJECT METHOD SPECIFICATION",
    "RESULTS",
    "DISCUSSION AND CONCLUSION",
}


def fix_heading_numbering_indents(doc) -> None:
    """Patch the multilevel-list definition in ``numbering.xml`` so that
    auto-numbered headings (``3.1``, ``3.2``, ``4.1`` …) sit flush with
    the left margin instead of being pushed one tab to the right.

    The original template used ``<w:ind w:left="X" w:hanging="X"/>`` and
    ``<w:suff w:val="tab"/>`` for every Heading-tied level, which makes
    Word render the heading as ``[number]<tab>[text]``, with the text
    aligned at column X. Setting ``left=0 hanging=0`` and replacing the
    tab suffix with a single space brings the text right next to the
    number, e.g. ``3.1 Methodological Framework``.
    """
    try:
        numbering_part = doc.part.numbering_part
    except (AttributeError, KeyError):
        return
    if numbering_part is None:
        return
    numbering_root = numbering_part.element

    patched = 0
    for lvl in numbering_root.iter(qn("w:lvl")):
        pStyle = lvl.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        style_val = pStyle.get(qn("w:val"), "")
        if not style_val.startswith("Heading"):
            continue

        # 1. Replace the "tab" suffix with a single space.
        suff = lvl.find(qn("w:suff"))
        if suff is None:
            suff = OxmlElement("w:suff")
            # Insert near the top of the level element.
            lvl.insert(0, suff)
        suff.set(qn("w:val"), "space")

        # 2. Zero out the indent inside <w:pPr>.
        ppr = lvl.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            lvl.append(ppr)
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:hanging"), "0")
        ind.set(qn("w:firstLine"), "0")
        # Drop any stray start/end attributes that override the zero.
        for stray in (qn("w:start"), qn("w:end")):
            if stray in ind.attrib:
                del ind.attrib[stray]

        patched += 1

    print(f"  patched {patched} heading-tied multilevel-list levels")


def reset_heading_indents(doc) -> None:
    """Force every Heading 1–4 style and paragraph to start flush with
    the left margin.

    Word's Multilevel List feature renders auto-numbered headings as
    ``<number><tab><heading text>`` and uses the heading style's left
    indent + first-line indent to position the number and the text. If
    the original template defined a non-zero left indent on Heading 2
    (which is common in many institutional templates) every numbered
    sub-heading then appears shifted one tab to the right of the page
    margin. The fix is to zero out the indent on the styles AND on each
    individual heading paragraph (because direct paragraph-level
    overrides win over style defaults).
    """
    target_styles = ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]

    for style_name in target_styles:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        pf = style.paragraph_format
        pf.left_indent = Inches(0)
        pf.first_line_indent = Inches(0)
        # Clear all tab stops on the style.
        with contextlib.suppress(Exception):
            pf.tab_stops.clear_all()
        # Also strip the underlying XML <w:ind> element so it doesn't
        # carry hidden hanging-indent values.
        ppr = style.element.find(qn("w:pPr"))
        if ppr is not None:
            ind = ppr.find(qn("w:ind"))
            if ind is not None:
                ppr.remove(ind)

    # Walk every Heading-styled paragraph and clear direct
    # paragraph-level indent overrides too.
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if not style_name.startswith("Heading"):
            continue
        ppr = p._p.find(qn("w:pPr"))
        if ppr is not None:
            ind = ppr.find(qn("w:ind"))
            if ind is not None:
                ppr.remove(ind)
        # Zero out via the python-docx API as well, for good measure.
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0)


def add_chapter_page_breaks(doc) -> None:
    """Insert a page break before every Heading 1 chapter title."""
    for p in doc.paragraphs:
        text = p.text.strip()
        if text not in CHAPTER_TITLES:
            continue
        style = p.style.name if p.style else ""
        if not style.startswith("Heading"):
            continue
        # Skip if the previous paragraph already ends with a page break.
        prev_xml = p._p.getprevious()
        if prev_xml is not None:
            br = prev_xml.find(qn("w:r") + "/" + qn("w:br"))
            if br is not None and br.get(qn("w:type")) == "page":
                continue

        # Prepend a page-break run to this paragraph.
        first_run = p.runs[0] if p.runs else p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        # Page break must come BEFORE the first text element.
        first_run._r.insert(0, br)


# ──────────────────────────────────────────────────────────────────────
#  Driver
# ──────────────────────────────────────────────────────────────────────

def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", default=DEFAULT_INPUT, type=Path)
    ap.add_argument("--output", default=DEFAULT_OUTPUT, type=Path)
    args = ap.parse_args()

    if not args.input.exists():
        print(f"[error] input file not found: {args.input}", file=sys.stderr)
        return 2

    print(f"[info] loading {args.input.name} …")
    doc = docx.Document(str(args.input))

    print("[info] setting document defaults (font, spacing, margins) …")
    set_document_defaults(doc)

    print("[info] reducing keywords to the five-keyword maximum …")
    reduce_keywords(doc)

    print("[info] fixing PDF→Word hyphenation artefacts …")
    fix_hyphenation_artefacts(doc)

    print("[info] sorting References alphabetically …")
    sort_references_alphabetically(doc)

    print("[info] inserting Table of Contents / List of Tables / List of Figures fields …")
    insert_toc_fields(doc)

    print("[info] applying Caption style to existing table and figure captions …")
    apply_caption_style(doc)

    print("[info] applying 11pt overrides for abstract, keywords, captions and references …")
    apply_typography_overrides(doc)

    print("[info] resetting heading indents (fixes 3.1, 3.2 tab-shifted bug) …")
    reset_heading_indents(doc)

    print("[info] patching multilevel-list (numbering.xml) for headings …")
    fix_heading_numbering_indents(doc)

    print("[info] adding page breaks before each chapter …")
    add_chapter_page_breaks(doc)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    print(f"[done] wrote {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
